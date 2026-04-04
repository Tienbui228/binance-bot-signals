#!/usr/bin/env python3
"""
build_daily_review_pack.py  —  Daily Review Report V2
======================================================
Drop-in replacement for V1. Called by BinanceScanner via subprocess.
CLI signature is IDENTICAL to V1 — no changes needed in oi_scanner.py.

V2 structure (all sections, phases 1-4):
  [COVER]
  A. Data Quality Gate
  B. Semantic Health Panel
  C. Confirm Path Integrity Board
  D. Executive Decision Summary
  E. Strategy Outcome Matrix
  F. Missed Opportunity Board
  G. Bad Loss Filter Board
  H. Decision Trace Coverage
  I. Human Review Queue
  J. Case Registry Preview        <- V1 preserved
  [page break per case]
  K. Case Detail (per-case)       <- V1 preserved verbatim
  L. Decision Footer

Locked semantic rules:
  - confirmed != sent (always distinct)
  - confirmed != close (confirmed_ts != closed_ts)
  - entry_or_confirm keyed to confirm time only
  - not_reached_yet is explicit, never substituted with fake timestamp
  - no strategy / regime / dispatch / veto changes
  - no semantic repair in renderer
  - semantically broken confirmed rows excluded from optimization counts
"""

from __future__ import annotations

import argparse
import csv
import json
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from zoneinfo import ZoneInfo


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
CLR_DARK_BLUE  = "1F4E79"
CLR_MID_BLUE   = "2E75B6"
CLR_LIGHT_BLUE = "D9EAF7"
CLR_GREEN_BG   = "E2EFDA"
CLR_RED_BG     = "FFE2E2"
CLR_AMBER_BG   = "FFF2CC"
CLR_ROW_ALT    = "F2F4F7"
CLR_WHITE      = "FFFFFF"
CLR_GREY_LABEL = "F2F4F7"
CLR_BORDER     = "BFBFBF"
CLR_GREEN_TXT  = "007600"
CLR_RED_TXT    = "CC0000"
CLR_AMBER_TXT  = "7B3F00"
CLR_GREY_TXT   = "666666"

VALID_3A_LABELS = {
    "trend_continuation_friendly",
    "broad_weakness_sell_pressure",
    "unclear_mixed",
}
LEGACY_LABELS = {
    "BULLISH", "BEARISH", "RANGE", "NEUTRAL", "UNKNOWN",
    "bullish", "bearish", "range", "neutral", "unknown",
}
VALID_FIT = {"HIGH", "MEDIUM", "LOW"}


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def read_csv_rows(path: Path) -> List[Dict]:
    if not path or not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def safe_float(v, default: float = 0.0) -> float:
    try:
        return float(v) if v not in (None, "", "NA", "N/A") else default
    except Exception:
        return default


def safe_int(v, default: int = 0) -> int:
    try:
        return int(float(v)) if v not in (None, "", "NA", "N/A") else default
    except Exception:
        return default


def yn(v) -> str:
    return "Y" if str(v or "").strip().upper() == "Y" else "N"


def fmt_ts(ts_ms: int, tz) -> str:
    if not ts_ms:
        return ""
    return (datetime.fromtimestamp(ts_ms / 1000.0, tz=timezone.utc)
            .astimezone(tz).strftime("%Y-%m-%d %H:%M:%S"))


def fmt_time(ts_ms: int, tz) -> str:
    if not ts_ms:
        return ""
    return (datetime.fromtimestamp(ts_ms / 1000.0, tz=timezone.utc)
            .astimezone(tz).strftime("%H:%M:%S"))


def owner_day(ts_ms: int, tz) -> str:
    if not ts_ms:
        return ""
    return (datetime.fromtimestamp(ts_ms / 1000.0, tz=timezone.utc)
            .astimezone(tz).strftime("%Y-%m-%d"))


def pct_str(n: int, total: int) -> str:
    return f"{n / total * 100:.0f}%" if total else "n/a"


def infer_close_type(row: Dict, fallback_hours: int) -> str:
    explicit = (row.get("case_close_type") or "").strip()
    if explicit:
        return explicit
    status = (row.get("status") or "").upper()
    if status in {"INVALIDATED", "EXPIRED_WAIT"}:
        return "true_close"
    created = safe_int(row.get("created_ts_ms"))
    if created:
        due = created + fallback_hours * 3600 * 1000
        now_ms = int(datetime.now(tz=timezone.utc).timestamp() * 1000)
        if status == "PENDING" and now_ms >= due:
            return "fallback_4h_snapshot"
    return "not_due_yet"


# ---------------------------------------------------------------------------
# Workspace / snapshot loading (V1-compatible)
# ---------------------------------------------------------------------------

def load_workspace_case(workspace: Path, case_day: str, case_id: str) -> Dict:
    if not workspace or not workspace.exists():
        return {}
    direct = workspace / "cases" / case_day / case_id / "case_meta.json"
    if direct.exists():
        try:
            return json.loads(direct.read_text(encoding="utf-8"))
        except Exception:
            pass
    for p in workspace.rglob("*.json"):
        if case_id in str(p):
            try:
                data = json.loads(p.read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    return data
            except Exception:
                pass
    return {}


def load_snapshot_index(path: Path) -> Dict[Tuple, Dict]:
    rows = read_csv_rows(path)
    out: Dict[Tuple, Dict] = {}
    for row in rows:
        pid   = row.get("pending_id") or row.get("setup_id") or ""
        stage = row.get("stage") or ""
        if pid and stage:
            out[(pid, stage)] = row
    return out


def infer_stage(row: Dict, case_meta: Dict, snapshots: Dict, stage_key: str) -> Tuple:
    pid = row.get("pending_id") or row.get("setup_id") or ""
    snap_key_map = {
        "pre_pending":    "pre_pending",
        "pending_open":   "pending",
        "entry_or_confirm": "confirmed",
        "case_close":     "closed",
    }
    snap_row = snapshots.get((pid, snap_key_map[stage_key]))
    img_path = None
    if snap_row and snap_row.get("image_path"):
        p = Path(snap_row["image_path"])
        if p.exists():
            img_path = p

    # try workspace case first
    stages = (case_meta.get("stages") or {}) if case_meta else {}
    item = stages.get(stage_key) or {}
    w_status = str(item.get("stage_status") or "")
    w_ctype  = str(item.get("stage_content_type") or "")
    w_note   = str(item.get("note") or "")
    if w_status:
        return w_status, w_ctype or ("chart_snapshot" if img_path else "none"), w_note, img_path

    # fallback inference
    final_status = (row.get("status") or "").upper()
    if stage_key in {"pre_pending", "pending_open"}:
        if img_path:
            return "captured", "chart_snapshot", "", img_path
        return "missing_unexpected", "none", "MISSING UNEXPECTED", None
    if stage_key == "entry_or_confirm":
        if img_path:
            return "captured", "chart_snapshot", "", img_path
        is_conf = (yn(row.get("is_confirmed")) == "Y"
                   or yn(row.get("is_sent_signal")) == "Y"
                   or final_status == "CONFIRMED")
        if is_conf:
            return "missing_unexpected", "none", "MISSING UNEXPECTED", None
        return "not_applicable", "none", "NOT APPLICABLE", None
    if stage_key == "case_close":
        if img_path:
            return "captured", "chart_snapshot", row.get("close_reason",""), img_path
        if infer_close_type(row, 4) == "not_due_yet":
            return "not_reached_yet", "none", "NOT REACHED YET", None
        return "missing_unexpected", "none", "MISSING UNEXPECTED", None
    return "", "", "", img_path


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def ensure_styles(doc: Document):
    s = doc.styles
    for name, base, size, bold in [
        ("CaseTitle",     "Heading 2",  13, True),
        ("MiniLabel",     "Body Text",   9, True),
        ("MetaText",      "Body Text",   9, False),
        ("CodeNote",      "Body Text",   9, False),
    ]:
        if name not in s:
            st = s.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = s[base]
            st.font.size  = Pt(size)
            st.font.bold  = bold
            if name == "CodeNote":
                st.font.name = "Courier New"


def _shd(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)


def _borders(cell, color: str = CLR_BORDER):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tb.append(b)
    tcPr.append(tb)


def _margins(cell, top=80, start=100, bottom=80, end=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    tcPr.append(m)
    for nm, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        n = OxmlElement(f"w:{nm}")
        n.set(qn("w:w"), str(val))
        n.set(qn("w:type"), "dxa")
        m.append(n)


def _cell(cell, text: str, bold: bool = False, size: int = 9,
          fg: str = "000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text or ""))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(fg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _repeat_hdr(row):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:tblHeader")
    h.set(qn("w:val"), "true")
    trPr.append(h)


def section_heading(doc: Document, letter: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    for txt, clr in [(f"{letter}.  {title}", CLR_DARK_BLUE)]:
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(clr)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), CLR_MID_BLUE)
    pBdr.append(bot)
    pPr.append(pBdr)


def info_line(doc: Document, label: str, value: str, size: int = 10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(size)
    r2 = p.add_run(str(value))
    r2.font.size = Pt(size)


def status_line(doc: Document, label: str, value: str, ok: bool):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    icon = "✓" if ok else "✗"
    clr  = CLR_GREEN_TXT if ok else CLR_RED_TXT
    r1 = p.add_run(f"{icon}  {label}: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor.from_string(clr)
    r2 = p.add_run(str(value))
    r2.font.size = Pt(10)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(clr)


def notice(doc: Document, text: str, kind: str = "warn"):
    clrs  = {"warn": CLR_AMBER_TXT, "pass": CLR_GREEN_TXT, "fail": CLR_RED_TXT}
    icons = {"warn": "⚠", "pass": "✓", "fail": "✗"}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.15)
    r = p.add_run(f"{icons.get(kind,'•')}  {text}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(clrs.get(kind, "000000"))


def make_table(doc: Document, headers: List[str], widths: List[float],
               hdr_bg: str = CLR_DARK_BLUE, hdr_fg: str = CLR_WHITE):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    _repeat_hdr(hdr)
    for i, (h, w) in enumerate(zip(headers, widths)):
        c = hdr.cells[i]
        c.width = Inches(w)
        _shd(c, hdr_bg)
        _borders(c)
        _margins(c)
        _cell(c, h, bold=True, fg=hdr_fg, align=WD_ALIGN_PARAGRAPH.CENTER)
    return tbl


def add_row(tbl, values: List[str], widths: List[float],
            alt: bool = False, bgs: Optional[List[str]] = None):
    row = tbl.add_row()
    default = CLR_ROW_ALT if alt else CLR_WHITE
    for i, (v, w) in enumerate(zip(values, widths)):
        c = row.cells[i]
        c.width = Inches(w)
        _shd(c, bgs[i] if bgs and i < len(bgs) else default)
        _borders(c)
        _margins(c)
        _cell(c, v, size=9)


# ---------------------------------------------------------------------------
# V1 helpers (preserved for case detail)
# ---------------------------------------------------------------------------

def add_kv_table(doc: Document, items: List[Tuple[str, str]], widths=(1.8, 4.9)):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    for k, v in items:
        r = tbl.add_row()
        r.cells[0].width = Inches(widths[0])
        r.cells[1].width = Inches(widths[1])
        _cell(r.cells[0], k, bold=True, size=9)
        _cell(r.cells[1], v, size=9)
        _shd(r.cells[0], CLR_GREY_LABEL)
        for c in r.cells:
            _margins(c)
    doc.add_paragraph("")


def add_stage_block(doc: Document, title: str, status: str, ctype: str,
                    note: str, img: Optional[Path], extras: Optional[List[str]] = None):
    doc.add_paragraph(title, style="CaseTitle")
    doc.add_paragraph(f"Stage status: {status or '-'} | Content type: {ctype or '-'}",
                      style="MetaText")
    for line in (extras or []):
        if line:
            doc.add_paragraph(line, style="MetaText")
    if note:
        doc.add_paragraph(f"Note: {note}", style="MetaText")
    if img and img.exists():
        try:
            doc.add_picture(str(img), width=Inches(6.3))
            cap = doc.add_paragraph(style="MetaText")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(img.name)
        except Exception as e:
            doc.add_paragraph(f"[image insert failed] {img.name}: {e}", style="CodeNote")
    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

class ReportData:
    def __init__(self, case_day: str, tz, fallback_hours: int,
                 owner_rows: List[Dict], precomputed: List[Tuple]):
        self.case_day       = case_day
        self.tz             = tz
        self.fallback_hours = fallback_hours
        self.owner_rows     = owner_rows
        self.precomputed    = precomputed
        self._compute()

    def _compute(self):
        rows = self.owner_rows
        self.total = len(rows)

        self.confirmed_rows = [r for r in rows if yn(r.get("is_confirmed")) == "Y"]
        self.sent_rows      = [r for r in rows if yn(r.get("is_sent_signal")) == "Y"]
        self.review_eligible = sum(1 for r in rows if yn(r.get("review_eligible")) == "Y")
        self.excluded        = self.total - self.review_eligible
        self.sem_broken      = sum(1 for r in rows if yn(r.get("semantic_consistency")) == "N")

        self.broken_confirmed = [
            r for r in self.confirmed_rows
            if not str(r.get("confirmed_ts_ms") or "").strip()
        ]
        self.clean_confirmed = len(self.confirmed_rows) - len(self.broken_confirmed)
        self.sent_no_ts      = sum(1 for r in self.sent_rows
                                   if not str(r.get("sent_ts_ms") or "").strip())
        self.closed_no_reason = sum(
            1 for r in rows
            if str(r.get("closed_ts_ms") or "").strip()
            and not str(r.get("close_reason") or "").strip()
        )
        self.semantic_health_ok = (
            len(self.broken_confirmed) == 0
            and self.sent_no_ts == 0
            and self.closed_no_reason == 0
        )

        labels = [r.get("regime_label","") for r in rows]
        self.regime_counts = dict(Counter(labels))
        self.legacy_count  = sum(1 for l in labels
                                 if l in LEGACY_LABELS
                                 or (l and l not in VALID_3A_LABELS))
        fits = [r.get("regime_fit_for_strategy","") for r in rows]
        self.fit_counts    = dict(Counter(fits))
        self.not_eval_fit  = sum(1 for v in fits if v not in VALID_FIT)

        self.by_strategy   = defaultdict(list)
        for r in rows:
            self.by_strategy[(r.get("strategy","legacy_5m_retest"),
                               r.get("side","UNKNOWN"))].append(r)

        self.wait_short_rows   = [r for r in rows if r.get("outcome_conclusion_code") == "WAIT_TOO_SHORT_CANDIDATE"]
        self.kill_correct_rows = [r for r in rows if r.get("outcome_conclusion_code") == "KILL_CORRECT"]
        self.no_move_rows      = [r for r in rows if r.get("outcome_conclusion_code") == "NO_MEANINGFUL_MOVE"]
        self.regret_valid_rows = [r for r in rows if yn(r.get("regret_valid_YN")) == "Y"]
        self.bad_loss_rows     = [
            r for r in rows
            if r.get("status") in ("INVALIDATED","EXPIRED_WAIT")
            and safe_float(r.get("future_2h_max_favor_pct")) > 1.5
            and r.get("outcome_conclusion_code") != "KILL_CORRECT"
        ]

        btc = [r.get("btc_regime","") for r in rows if r.get("btc_regime")]
        mkt = [r.get("market_regime","") for r in rows if r.get("market_regime")]
        self.btc_regime_mode = Counter(btc).most_common(1)[0][0] if btc else "unknown"
        self.mkt_regime_mode = Counter(mkt).most_common(1)[0][0] if mkt else "unknown"

        self.missing_pre = self.missing_pend = self.missing_ent = self.missing_clo = 0
        for _, _, pre, pend, ent, clo in self.precomputed:
            if pre[0]  == "missing_unexpected": self.missing_pre  += 1
            if pend[0] == "missing_unexpected": self.missing_pend += 1
            if ent[0]  == "missing_unexpected": self.missing_ent  += 1
            if clo[0]  == "missing_unexpected": self.missing_clo  += 1

    def pct(self, n): return pct_str(n, self.total)

    def next_action(self) -> str:
        if not self.semantic_health_ok:
            return ("Fix semantic health first — broken confirmed-path rows "
                    "make strategy conclusions unreliable.")
        if self.legacy_count > 0:
            return (f"Normalize regime labels — {self.legacy_count} row(s) still carry "
                    "legacy labels. Sprint 3A.2 may have partial coverage gaps.")
        if len(self.wait_short_rows) >= 3:
            top = Counter((r.get("strategy","?"), r.get("side","?"))
                          for r in self.wait_short_rows).most_common(1)
            if top:
                fam, side = top[0][0]
                return (f"Review invalidation window for {fam}/{side} — "
                        f"{len(self.wait_short_rows)} WAIT_TOO_SHORT_CANDIDATE cases.")
        if self.not_eval_fit > self.total * 0.1:
            return (f"Improve regime_fit coverage — {self.not_eval_fit} rows "
                    "still have fit=not_evaluated.")
        if len(self.bad_loss_rows) >= 2:
            return (f"Inspect {len(self.bad_loss_rows)} bad-loss cases — market moved "
                    "strongly in their favour after kill.")
        return "Data clean — no single highest-priority action today."

    def caveats(self) -> List[str]:
        out = []
        if self.broken_confirmed:
            out.append(f"{len(self.broken_confirmed)} confirmed row(s) have semantic issues — excluded from optimization counts.")
        if self.legacy_count:
            out.append(f"{self.legacy_count} row(s) carry legacy regime labels.")
        if self.not_eval_fit:
            out.append(f"{self.not_eval_fit} row(s) have regime_fit=not_evaluated.")
        if not self.total:
            out.append("No pending rows found for this date partition.")
        return out


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def s_cover(doc, data: ReportData, generated_at: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Daily Signal Review Report V2")
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(CLR_DARK_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Binance Futures Signal Bot  ·  Manual Trading Review")
    r2.font.size = Pt(11); r2.italic = True
    r2.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)
    p2.paragraph_format.space_after = Pt(2)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Date: {data.case_day}   |   Generated: {generated_at}")
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)
    p3.paragraph_format.space_after = Pt(14)

    ok = data.semantic_health_ok
    health = ("CLEAN — safe for optimization" if ok
              else "⚠  ISSUES FOUND — check Semantic Health panel before conclusions")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"Health status: {health}")
    r4.bold = True; r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor.from_string(CLR_GREEN_TXT if ok else CLR_RED_TXT)
    p4.paragraph_format.space_after = Pt(14)
    doc.add_paragraph()


def s_a(doc, data: ReportData):
    section_heading(doc, "A", "Data Quality Gate")
    info_line(doc, "Report date", data.case_day)
    info_line(doc, "Total pending rows", str(data.total))
    info_line(doc, "Status breakdown",
              "  |  ".join(f"{k}={v}" for k, v in
                           sorted(Counter(r.get("status","") for r in data.owner_rows).items())))
    doc.add_paragraph()

    tbl = make_table(doc, ["Metric","Count","Of total","Status"],
                     [2.8, 0.8, 0.8, 1.0])
    rows_def = [
        ("Confirmed (is_confirmed=Y)", len(data.confirmed_rows), True),
        ("Sent (is_sent_signal=Y)",    len(data.sent_rows),      True),
        ("Clean confirmed",            data.clean_confirmed,     len(data.broken_confirmed)==0),
        ("Broken confirmed (excluded)",len(data.broken_confirmed),len(data.broken_confirmed)==0),
        ("Outcome 2h available",
         sum(1 for r in data.owner_rows if yn(r.get("outcome_2h_available"))=="Y"), True),
        ("Timestamps trustworthy",
         "Y" if not data.sem_broken else "N", not data.sem_broken),
    ]
    for i, (metric, count, ok) in enumerate(rows_def):
        bgs = ([CLR_ROW_ALT if i%2 else CLR_WHITE]*3
               + [CLR_GREEN_BG if ok else CLR_RED_BG])
        add_row(tbl, [metric, str(count), data.pct(count) if isinstance(count,int) else "",
                      "OK" if ok else "ISSUE"], [2.8,0.8,0.8,1.0], bgs=bgs)
    doc.add_paragraph()


def s_b(doc, data: ReportData):
    section_heading(doc, "B", "Semantic Health Panel")
    if data.semantic_health_ok:
        notice(doc, "All confirmed-path rows are semantically consistent.", "pass")
    else:
        notice(doc, "Semantic issues found — do not use optimization counts from broken rows.", "fail")
    doc.add_paragraph()
    for label, val, ok in [
        ("Confirmed rows with confirmed_ts_ms",    data.clean_confirmed,            data.clean_confirmed==len(data.confirmed_rows)),
        ("Confirmed rows missing confirmed_ts_ms", len(data.broken_confirmed),      len(data.broken_confirmed)==0),
        ("Sent rows missing sent_ts_ms",           data.sent_no_ts,                 data.sent_no_ts==0),
        ("Closed rows missing close_reason",       data.closed_no_reason,           data.closed_no_reason==0),
        ("Rows with semantic_consistency=N",       data.sem_broken,                 data.sem_broken==0),
        ("missing_unexpected: pre_pending",        data.missing_pre,                data.missing_pre==0),
        ("missing_unexpected: pending_open",       data.missing_pend,               data.missing_pend==0),
        ("missing_unexpected: entry_or_confirm",   data.missing_ent,                data.missing_ent==0),
        ("missing_unexpected: case_close",         data.missing_clo,                data.missing_clo==0),
    ]:
        status_line(doc, label, str(val), ok)
    doc.add_paragraph()


def s_c(doc, data: ReportData):
    section_heading(doc, "C", "Confirm Path Integrity Board")
    p = doc.add_paragraph()
    r = p.add_run("'not_reached_yet' is explicit — never substituted with a fake timestamp.")
    r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)
    p.paragraph_format.space_after = Pt(6)

    conf_rows = [r for r in data.owner_rows if yn(r.get("is_confirmed"))=="Y"]
    if not conf_rows:
        doc.add_paragraph("No confirmed rows for this date.")
        doc.add_paragraph(); return

    tbl = make_table(doc, ["Symbol","Side","confirmed_at","sent_at","closed_at","Integrity"],
                     [1.1, 0.6, 1.4, 1.4, 1.4, 1.6])
    tz = data.tz
    for i, row in enumerate(conf_rows[:25]):
        conf_ts   = safe_int(row.get("confirmed_ts_ms"))
        sent_ts   = safe_int(row.get("sent_ts_ms"))
        closed_ts = safe_int(row.get("closed_ts_ms"))
        # RULE: no fake timestamp — show "missing" or "not_reached_yet" explicitly
        c_disp = fmt_time(conf_ts, tz)  if conf_ts   else "missing"
        s_disp = fmt_time(sent_ts, tz)  if sent_ts   else "not_sent"
        x_disp = fmt_time(closed_ts,tz) if closed_ts else "not_reached_yet"
        ok = yn(row.get("semantic_consistency"))=="Y"
        bgs = [CLR_ROW_ALT if i%2 else CLR_WHITE]*5 + [CLR_GREEN_BG if ok else CLR_RED_BG]
        add_row(tbl, [row.get("symbol",""), row.get("side",""), c_disp, s_disp, x_disp,
                      "OK" if ok else f"BROKEN: {row.get('semantic_issue','')}"],
                [1.1,0.6,1.4,1.4,1.4,1.6], bgs=bgs)
    if len(conf_rows) > 25:
        doc.add_paragraph(f"  … and {len(conf_rows)-25} more confirmed rows.")
    doc.add_paragraph()


def s_d(doc, data: ReportData):
    section_heading(doc, "D", "Executive Decision Summary")
    info_line(doc, "BTC regime (mode)",  data.btc_regime_mode)
    info_line(doc, "Market regime (mode)",data.mkt_regime_mode)
    info_line(doc, "Cases detected today", str(data.total))
    info_line(doc, "Confirmed path cases", str(len(data.confirmed_rows)))
    info_line(doc, "Sent signals",         str(len(data.sent_rows)))

    long_tot  = sum(1 for r in data.owner_rows if r.get("side")=="LONG")
    short_tot = sum(1 for r in data.owner_rows if r.get("side")=="SHORT")
    long_conf  = sum(1 for r in data.confirmed_rows if r.get("side")=="LONG")
    short_conf = sum(1 for r in data.confirmed_rows if r.get("side")=="SHORT")
    info_line(doc, "Long detected / confirmed",
              f"{long_tot} / {long_conf} ({pct_str(long_conf,long_tot)})")
    info_line(doc, "Short detected / confirmed",
              f"{short_tot} / {short_conf} ({pct_str(short_conf,short_tot)})")
    info_line(doc, "WAIT_TOO_SHORT_CANDIDATE", str(len(data.wait_short_rows)))
    info_line(doc, "KILL_CORRECT",             str(len(data.kill_correct_rows)))
    info_line(doc, "Regret-valid (regret_valid=Y)", str(len(data.regret_valid_rows)))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run("→  Next action: ")
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor.from_string(CLR_MID_BLUE)
    r2 = p.add_run(data.next_action())
    r2.bold = True; r2.font.size = Pt(11)
    doc.add_paragraph()


def s_e(doc, data: ReportData):
    section_heading(doc, "E", "Strategy Outcome Matrix")
    p = doc.add_paragraph()
    r = p.add_run("Truth-clean subset only. Broken confirmed rows excluded.")
    r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)
    p.paragraph_format.space_after = Pt(6)

    broken_ids = {r.get("pending_id") for r in data.broken_confirmed}
    tbl = make_table(doc,
        ["Strategy | Side","Det.","Conf.","Sent","KILL_OK","WAIT_SHORT","NO_MOVE","avg Score"],
        [2.1, 0.5, 0.5, 0.5, 0.7, 0.9, 0.7, 0.8])
    for i, ((strat,side), rows) in enumerate(sorted(data.by_strategy.items())):
        clean = [r for r in rows if r.get("pending_id") not in broken_ids]
        scores = [safe_float(r.get("score")) for r in rows if safe_float(r.get("score"))>0]
        add_row(tbl, [
            f"{strat} | {side}", str(len(rows)),
            str(sum(1 for r in clean if yn(r.get("is_confirmed"))=="Y")),
            str(sum(1 for r in clean if yn(r.get("is_sent_signal"))=="Y")),
            str(sum(1 for r in clean if r.get("outcome_conclusion_code")=="KILL_CORRECT")),
            str(sum(1 for r in clean if r.get("outcome_conclusion_code")=="WAIT_TOO_SHORT_CANDIDATE")),
            str(sum(1 for r in clean if r.get("outcome_conclusion_code")=="NO_MEANINGFUL_MOVE")),
            f"{sum(scores)/len(scores):.1f}" if scores else "n/a",
        ], [2.1,0.5,0.5,0.5,0.7,0.9,0.7,0.8], alt=(i%2==1))
    doc.add_paragraph()


def s_f(doc, data: ReportData):
    section_heading(doc, "F", "Missed Opportunity Board")
    p = doc.add_paragraph()
    r = p.add_run("Cases bot killed/expired but market moved strongly in expected direction afterwards (WAIT_TOO_SHORT_CANDIDATE).")
    r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)
    p.paragraph_format.space_after = Pt(6)

    if not data.wait_short_rows:
        doc.add_paragraph("No WAIT_TOO_SHORT_CANDIDATE cases for this date.")
        doc.add_paragraph(); return

    tbl = make_table(doc, ["Symbol","Side","Strategy","Status","Fav 2h%","Adv 2h%","Close reason"],
                     [1.1, 0.6, 1.8, 1.1, 0.8, 0.8, 1.3])
    for i, row in enumerate(sorted(data.wait_short_rows,
                                   key=lambda r: safe_float(r.get("future_2h_max_favor_pct")),
                                   reverse=True)):
        add_row(tbl, [
            row.get("symbol",""), row.get("side",""), row.get("strategy",""),
            row.get("status",""),
            f"{safe_float(row.get('future_2h_max_favor_pct')):.2f}",
            f"{safe_float(row.get('future_2h_max_adverse_pct')):.2f}",
            row.get("close_reason",""),
        ], [1.1,0.6,1.8,1.1,0.8,0.8,1.3], alt=(i%2==1))
    doc.add_paragraph()


def s_g(doc, data: ReportData):
    section_heading(doc, "G", "Bad Loss Filter Board")
    p = doc.add_paragraph()
    r = p.add_run("KILL_CORRECT = kill justified. Bad loss = killed but market moved >1.5% in favour after kill.")
    r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)
    p.paragraph_format.space_after = Pt(6)

    info_line(doc, "KILL_CORRECT", str(len(data.kill_correct_rows)))
    info_line(doc, "Bad loss (fav >1.5% after kill)", str(len(data.bad_loss_rows)))
    info_line(doc, "NO_MEANINGFUL_MOVE", str(len(data.no_move_rows)))
    doc.add_paragraph()

    if data.kill_correct_rows:
        p2 = doc.add_paragraph()
        p2.add_run("KILL_CORRECT by close reason:").bold = True
        p2.runs[0].font.size = Pt(10)
        p2.paragraph_format.space_after = Pt(4)
        for reason, cnt in Counter(r.get("close_reason","unknown")
                                   for r in data.kill_correct_rows).most_common():
            info_line(doc, f"  {reason}", str(cnt))
        doc.add_paragraph()

    if data.bad_loss_rows:
        p3 = doc.add_paragraph()
        p3.add_run("Bad loss cases — inspect for rule improvement:").bold = True
        p3.runs[0].font.size = Pt(10)
        p3.paragraph_format.space_after = Pt(4)
        tbl = make_table(doc, ["Symbol","Side","Strategy","Fav 2h%","Adv 2h%","Close reason"],
                         [1.1, 0.6, 1.8, 0.8, 0.8, 2.4])
        for i, row in enumerate(sorted(data.bad_loss_rows,
                                       key=lambda r: safe_float(r.get("future_2h_max_favor_pct")),
                                       reverse=True)):
            add_row(tbl, [
                row.get("symbol",""), row.get("side",""), row.get("strategy",""),
                f"{safe_float(row.get('future_2h_max_favor_pct')):.2f}",
                f"{safe_float(row.get('future_2h_max_adverse_pct')):.2f}",
                row.get("close_reason",""),
            ], [1.1,0.6,1.8,0.8,0.8,2.4], alt=(i%2==1))
        doc.add_paragraph()


def s_h(doc, data: ReportData):
    section_heading(doc, "H", "Decision Trace Coverage")

    regime_ok   = sum(1 for r in data.owner_rows if r.get("regime_label","") in VALID_3A_LABELS)
    fit_ok      = sum(1 for r in data.owner_rows if r.get("regime_fit_for_strategy","") in VALID_FIT)
    dispatch_ok = sum(1 for r in data.owner_rows if r.get("dispatch_action","") not in ("not_evaluated","",None))
    setup_ok    = sum(1 for r in data.owner_rows if r.get("setup_quality_band","") not in ("not_evaluated","",None))

    tbl = make_table(doc, ["Field","Populated","Coverage","Status"],
                     [2.4, 0.9, 0.9, 1.0])
    for i, (field, pop, ok) in enumerate([
        ("regime_label (3A set)",    regime_ok,   data.legacy_count==0),
        ("regime_fit_for_strategy",  fit_ok,      data.not_eval_fit==0),
        ("dispatch_action",          dispatch_ok, True),
        ("setup_quality_band",       setup_ok,    True),
    ]):
        bgs = [CLR_ROW_ALT if i%2 else CLR_WHITE]*3 + [CLR_GREEN_BG if ok else CLR_AMBER_BG]
        add_row(tbl, [field, str(pop), data.pct(pop), "OK" if ok else "GAPS"],
                [2.4,0.9,0.9,1.0], bgs=bgs)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Regime label distribution:").bold = True
    p.runs[0].font.size = Pt(10); p.paragraph_format.space_after = Pt(4)
    for label, cnt in sorted(data.regime_counts.items(), key=lambda x: -x[1]):
        marker = "✓" if label in VALID_3A_LABELS else "✗ LEGACY"
        info_line(doc, f"  {label}  {marker}", str(cnt))

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run("Fit distribution:").bold = True
    p2.runs[0].font.size = Pt(10); p2.paragraph_format.space_after = Pt(4)
    for val, cnt in sorted(data.fit_counts.items(), key=lambda x: -x[1]):
        info_line(doc, f"  {val}", str(cnt))
    doc.add_paragraph()


def s_i(doc, data: ReportData):
    section_heading(doc, "I", "Human Review Queue")
    p = doc.add_paragraph()
    r = p.add_run("Priority: regret_valid=Y first, then outcome available, then score desc.")
    r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)
    p.paragraph_format.space_after = Pt(6)

    candidates = sorted(
        [r for r in data.owner_rows
         if r.get("status") in ("INVALIDATED","EXPIRED_WAIT","CONFIRMED")],
        key=lambda r: (
            yn(r.get("regret_valid_YN"))=="Y",
            yn(r.get("outcome_2h_available"))=="Y",
            safe_float(r.get("score")),
        ),
        reverse=True
    )[:20]

    if not candidates:
        doc.add_paragraph("No review candidates for this date.")
        doc.add_paragraph(); return

    tbl = make_table(doc,
        ["Symbol","Side","Strategy","Status","Score","Regret?","Outcome code"],
        [1.1, 0.6, 1.8, 1.1, 0.6, 0.7, 1.6])
    for i, row in enumerate(candidates):
        add_row(tbl, [
            row.get("symbol",""), row.get("side",""), row.get("strategy",""),
            row.get("status",""),
            f"{safe_float(row.get('score')):.1f}",
            row.get("regret_valid_YN",""),
            row.get("outcome_conclusion_code",""),
        ], [1.1,0.6,1.8,1.1,0.6,0.7,1.6], alt=(i%2==1))
    doc.add_paragraph()


def s_j(doc, data: ReportData):
    section_heading(doc, "J", "Case Registry Preview")
    tbl = make_table(doc, ["Case ID","Symbol","Side","Strategy","Status","Outcome"],
                     [1.5, 1.0, 0.6, 1.8, 1.0, 1.6])
    for row, _, _, _, _, _ in data.precomputed:
        add_row(tbl, [
            row.get("pending_id") or row.get("setup_id",""),
            row.get("symbol",""), row.get("side",""), row.get("strategy",""),
            row.get("status",""),
            row.get("outcome_conclusion_code","") or row.get("close_reason",""),
        ], [1.5,1.0,0.6,1.8,1.0,1.6])
    doc.add_paragraph()


def s_k(doc, data: ReportData):
    tz = data.tz
    fh = data.fallback_hours
    for idx, (row, _, pre, pend, ent, clo) in enumerate(data.precomputed, 1):
        doc.add_page_break()
        case_id = row.get("pending_id") or row.get("setup_id","")
        doc.add_paragraph(f"Case {idx}: {case_id}", style="CaseTitle")

        cts = safe_int(row.get("created_ts_ms"))
        add_kv_table(doc, [
            ("Summary",
             f"{row.get('symbol','')} | {row.get('side','')} | {row.get('strategy','')} | "
             f"status={row.get('status','')} | close_reason={row.get('close_reason','')}"),
            ("Times",
             f"signal={fmt_ts(safe_int(row.get('signal_open_time')),tz)} | "
             f"created={fmt_ts(cts,tz)} | "
             f"confirmed={fmt_ts(safe_int(row.get('confirmed_ts_ms')),tz)} | "
             f"sent={fmt_ts(safe_int(row.get('sent_ts_ms')),tz)} | "
             f"close={fmt_ts(safe_int(row.get('closed_ts_ms')),tz)} | "
             f"fallback_due={fmt_ts(cts+fh*3600000,tz) if cts else ''}"),
            ("Semantics",
             f"is_confirmed={yn(row.get('is_confirmed'))} | "
             f"is_sent_signal={yn(row.get('is_sent_signal'))} | "
             f"review_eligible={yn(row.get('review_eligible'))} | "
             f"semantic_consistency={yn(row.get('semantic_consistency'))}"),
            ("Review integrity",
             f"review_exclusion_reason={row.get('review_exclusion_reason','')} | "
             f"semantic_issue={row.get('semantic_issue','')}"),
            ("Close anchoring",
             f"case_close_type={infer_close_type(row,fh)} | "
             f"close_anchor_time={fmt_ts(safe_int(row.get('close_anchor_time_ms')),tz)} | "
             f"close_capture_basis={row.get('close_capture_basis','')}"),
            ("Regime / dispatch",
             f"regime_label={row.get('regime_label','')} | "
             f"fit={row.get('regime_fit_for_strategy','')} | "
             f"dispatch={row.get('dispatch_action','')} [{row.get('dispatch_confidence_band','')}]"),
            ("Tradeability",
             f"entry_feasible={row.get('entry_feasible_YN','')} | "
             f"slippage={row.get('entry_slippage_pct','')} | "
             f"regret_valid={row.get('regret_valid_YN','')} | "
             f"outcome={row.get('outcome_conclusion_code','')}"),
        ])

        add_stage_block(doc, "Pre Pending",      pre[0],  pre[1],  pre[2],  pre[3])
        add_stage_block(doc, "Pending Open",     pend[0], pend[1], pend[2], pend[3])
        add_stage_block(doc, "Entry Or Confirm", ent[0],  ent[1],  ent[2],  ent[3],
                        [f"confirm_fail_detail: {row.get('confirm_fail_detail','')}"
                         ] if row.get("confirm_fail_detail") else None)
        add_stage_block(doc, "Case Close",       clo[0],  clo[1],  clo[2],  clo[3],
                        [l for l in [
                            f"close_trigger_detail: {row.get('close_trigger_detail','')}" if row.get("close_trigger_detail") else "",
                            f"invalidation_detail: {row.get('invalidation_detail','')}"    if row.get("invalidation_detail")    else "",
                        ] if l])

        doc.add_paragraph("Post-Close Outcome", style="CaseTitle")
        add_kv_table(doc, [
            ("1h outcome",
             f"available={yn(row.get('outcome_1h_available'))} | "
             f"favor={row.get('future_1h_max_favor_pct','')} | "
             f"adverse={row.get('future_1h_max_adverse_pct','')}"),
            ("2h outcome",
             f"available={yn(row.get('outcome_2h_available'))} | "
             f"favor={row.get('future_2h_max_favor_pct','')} | "
             f"adverse={row.get('future_2h_max_adverse_pct','')}"),
            ("4h reference",
             f"favor={row.get('future_4h_max_favor_pct','')} | "
             f"adverse={row.get('future_4h_max_adverse_pct','')}"),
            ("Reclaim",
             f"2h={row.get('reclaim_breakout_2h_YN','')} | "
             f"4h={row.get('reclaim_breakout_4h_YN','')}"),
            ("Execution",
             f"entry_feasible={row.get('entry_feasible_YN','')} | "
             f"window={row.get('entry_feasible_window_minutes','')}m | "
             f"note={row.get('entry_execution_note','')}"),
            ("Time-to-move",
             f"to_favor={row.get('time_to_max_favor_minutes','')}m | "
             f"to_adverse={row.get('time_to_max_adverse_minutes','')}m"),
            ("Regret filter",
             f"regret_valid={row.get('regret_valid_YN','')} | "
             f"reason={row.get('regret_filter_reason','')}"),
            ("Conclusion",
             f"{row.get('outcome_conclusion_code','')} | "
             f"notes={row.get('post_close_outcome_notes','')}"),
        ])
        doc.add_paragraph("Review placeholders", style="CaseTitle")
        doc.add_paragraph(
            "human_review_status=PENDING | verdict_code= | root_cause_code= | action_candidate_code=",
            style="MetaText",
        )


def s_l(doc, data: ReportData):
    doc.add_page_break()
    section_heading(doc, "L", "Decision Footer")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r1 = p.add_run("→  Next action: ")
    r1.bold = True; r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor.from_string(CLR_MID_BLUE)
    r2 = p.add_run(data.next_action())
    r2.bold = True; r2.font.size = Pt(12)

    caveats = data.caveats()
    if caveats:
        doc.add_paragraph("Caveats:").runs[0].bold = True
        for c in caveats:
            notice(doc, c, "warn")
    doc.add_paragraph()

    p3 = doc.add_paragraph()
    r3 = p3.add_run("Semantic rules locked in this report:")
    r3.bold = True; r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)
    p3.paragraph_format.space_after = Pt(2)
    for rule in [
        "confirmed ≠ sent (always distinct)",
        "confirmed ≠ close (confirmed_ts ≠ closed_ts)",
        "entry_or_confirm keyed to confirm time only",
        "not_reached_yet is explicit — never substituted with fake timestamp",
        "no strategy / regime / dispatch / veto changes in this report",
        "semantically broken confirmed rows excluded from optimization counts",
    ]:
        pr = doc.add_paragraph()
        pr.paragraph_format.left_indent = Inches(0.2)
        pr.paragraph_format.space_after = Pt(1)
        rr = pr.add_run(f"• {rule}")
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)

    doc.add_paragraph()
    psig = doc.add_paragraph()
    psig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsig = psig.add_run(
        f"Daily Review Report V2  |  {data.case_day}  |  "
        f"cases={data.total}  |  confirmed={len(data.confirmed_rows)}  |  "
        f"sent={len(data.sent_rows)}"
    )
    rsig.font.size = Pt(9); rsig.italic = True
    rsig.font.color.rgb = RGBColor.from_string(CLR_GREY_TXT)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build_doc(args):
    tz           = ZoneInfo(args.tz)
    fallback_hrs = int(args.fallback_hours)
    workspace    = Path(args.workspace)
    snap_idx     = load_snapshot_index(Path(args.snapshot_index)) if args.snapshot_index else {}

    all_rows = read_csv_rows(Path(args.pending))
    owner_rows = sorted(
        [r for r in all_rows
         if owner_day(safe_int(r.get("created_ts_ms")), tz) == args.date],
        key=lambda r: safe_int(r.get("created_ts_ms"))
    )

    precomputed = []
    for row in owner_rows:
        cid = row.get("pending_id") or row.get("setup_id","")
        cm  = load_workspace_case(workspace, args.date, cid)
        precomputed.append((
            row, cm,
            infer_stage(row, cm, snap_idx, "pre_pending"),
            infer_stage(row, cm, snap_idx, "pending_open"),
            infer_stage(row, cm, snap_idx, "entry_or_confirm"),
            infer_stage(row, cm, snap_idx, "case_close"),
        ))

    data = ReportData(args.date, tz, fallback_hrs, owner_rows, precomputed)
    generated_at = (datetime.now(tz=timezone.utc).astimezone(tz)
                    .strftime("%Y-%m-%d %H:%M:%S"))

    doc = Document()
    ensure_styles(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.8)
    sec.top_margin  = sec.bottom_margin = Inches(0.7)

    s_cover(doc, data, generated_at)
    s_a(doc, data)
    s_b(doc, data)
    s_c(doc, data)
    s_d(doc, data)
    s_e(doc, data)
    s_f(doc, data)
    s_g(doc, data)
    s_h(doc, data)
    s_i(doc, data)
    s_j(doc, data)
    s_k(doc, data)
    s_l(doc, data)

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"daily_review_{args.date}.docx"
    doc.save(str(out_path))
    print(f"[review_case builder] wrote {out_path}")


def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument("--date",           required=True)
    p.add_argument("--workspace",      required=True)
    p.add_argument("--pending",        required=True)
    p.add_argument("--signals",        required=False, default="")
    p.add_argument("--results",        required=False, default="")
    p.add_argument("--snapshot-index", required=False, default="")
    p.add_argument("--out-dir",        required=True)
    p.add_argument("--tz",             required=False, default="Asia/Ho_Chi_Minh")
    p.add_argument("--fallback-hours", required=False, default="4")
    return p.parse_args()


if __name__ == "__main__":
    build_doc(parse_args())

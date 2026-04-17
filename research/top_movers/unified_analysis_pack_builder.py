"""
research/top_movers/unified_analysis_pack_builder.py

Builds R1_unified_analysis_pack_<research_day>.docx

A single consolidated DOCX containing all CSV/ledger data in readable form,
optimized for ChatGPT context ingestion and cross-day analysis.

7 sections:
  1. Scope & Interpretation Guide
  2. Multi-Day Executive Summary
  3. Daily Intervention Snapshot
  4. Daily Signature Candidates
  5. Cross-Day Ledger Snapshot
  6. Raw Ledger Appendix
  7. Raw Daily Summary Appendix

Downstream-only. Does NOT touch live runtime, lifecycle, or strategy files.
"""

import csv
import os
from datetime import datetime, timezone
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from research.top_movers.io import OUTPUT_BASE
from research.top_movers.signature_ledger import (
    LEDGER_PATH,
    load_and_normalize_ledger_rows,
    ledger_rows_as_of,
    build_ledger_snapshot_for_report,
    build_intervention_shortlist,
    # Phase 2B
    build_multiday_family_stats,
    build_controlled_validation_summary,
    build_multiday_interaction_stats,
    # Phase 2D
    build_family_validation_stats,
    build_family_answer_contracts,
)

# ---------------------------------------------------------------------------
# Docx helpers — self-contained, not imported from docx_report_builder
# to avoid coupling between the two render owners.
# ---------------------------------------------------------------------------

def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _white_bold(cell) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _font(cell, bold: bool = False, size: int = 8) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold
            r.font.size = Pt(size)


def _make_table(doc, headers: List[str], col_w: List[float], fill: str = "2E75B6"):
    if sum(col_w) > 6.5: col_w = [w * 6.5 / sum(col_w) for w in col_w]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(sum(col_w) * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for i, cell in enumerate(t.rows[0].cells):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(col_w[i] * 1440)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.insert(0, tcW)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        _shade(cell, fill)
        _white_bold(cell)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t


def _add_row(t, vals: List, hi_col: int = -1, hi_map: Optional[Dict] = None) -> None:
    row = t.add_row()
    for i, v in enumerate(vals):
        cell = row.cells[i]
        cell.text = str(v) if v is not None else "—"
        _font(cell)
        if hi_map and i == hi_col:
            color = hi_map.get(str(v))
            if color:
                _shade(cell, color)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _h1(doc, text: str) -> None:
    doc.add_heading(text, level=1)


def _h2(doc, text: str) -> None:
    doc.add_heading(text, level=2)


def _p(doc, text: str, bold: bool = False, size: int = 10, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.italic = italic
    return p


def _fmt(v, nd: int = 3) -> str:
    if v is None or v == "" or v == "None":
        return "—"
    try:
        return str(round(float(v), nd))
    except Exception:
        return str(v)


# ---------------------------------------------------------------------------
# Color maps
# ---------------------------------------------------------------------------

_ROLE_COLORS = {
    "repeated_candidate": "70AD47",
    "tracking":           "FFC000",
    "first_observation":  "2E75B6",
    "stale":              "808080",
}

_HEALTH_COLORS_UA = {
    "CLEAN":                  "70AD47",
    "CLEAN_WITH_VISUAL_GAPS": "92D050",
    "PARTIAL":                "FFC000",
    "WEAK":                   "C00000",
}

_GRADE_COLORS = {
    "OLD_STRATEGY_IMPROVEMENT_CANDIDATE": "2E75B6",
    "NEW_STRATEGY_THESIS_CANDIDATE":      "70AD47",
    "KEEP_TRACKING":                      "FFC000",
    "DESCRIPTIVE_ONLY":                   "808080",
    "NOT_RELIABLE_YET":                   "C00000",
}

# Phase 2B color maps
_PROMOTION_COLORS_UA = {
    "DESCRIPTIVE_ONLY":                    "808080",
    "KEEP_TRACKING":                       "FFC000",
    "PREPARE_HYPOTHESIS":                  "2E75B6",
    "NOT_READY_FOR_CONTROLLED_VALIDATION": "C00000",
}
_SAMPLE_GATE_COLORS_UA = {
    "RECOMMENDATION_GRADE": "70AD47",
    "TRACKING_GRADE":       "FFC000",
    "LOW_SAMPLE":           "C0504D",
    "NOT_ENOUGH_SAMPLE":    "808080",
}
_CONFIDENCE_2B_COLORS_UA = {
    "MODERATE":         "FFC000",
    "LOW":              "C00000",
    "DESCRIPTIVE_ONLY": "808080",
}


# ---------------------------------------------------------------------------
# Section 1 — Scope & Interpretation Guide
# ---------------------------------------------------------------------------

def _u1_interpretation_guide(doc) -> None:
    _h2(doc, "1. Scope & Interpretation Guide")
    _p(doc, "What this file is:", bold=True)
    _p(doc,
       "This is the R1 Unified Analysis Pack — a consolidated DOCX containing all daily "
       "CSV and ledger evidence in one artifact, optimized for cross-day analysis "
       "and ChatGPT context ingestion. It is downstream-only and does not affect live "
       "strategy, runtime, or lifecycle behavior.",
       size=9)

    _p(doc, "Critical distinctions (read before interpreting any section):", bold=True)
    t = _make_table(doc, ["Rule", "Explanation"], [2.3, 6.1], fill="404040")

    rules = [
        ("Repeated signatures ≠ strategy proof",
         "A repeated signature means the same (side, pre_move_signature, participation_pattern, "
         "structural_quality) combo appeared in ≥2 eligible cases on a given day. This is a "
         "pattern worth tracking — NOT a validated strategy or a recommendation to change live rules."),
        ("Case-level theses ≠ repeated patterns",
         "A case can have decision_grade = NEW_STRATEGY_THESIS_CANDIDATE or "
         "OLD_STRATEGY_IMPROVEMENT_CANDIDATE even when zero repeated signatures exist. "
         "These are per-case conclusions derived from individual case interpretation. "
         "They are independent of the signature evidence board."),
        ("Readiness levels",
         "descriptive_only: not enough repetition or strategy relevance | "
         "keep_tracking: interesting but insufficient for intervention | "
         "old_strategy_improvement_candidate: points to existing strategy + identifiable layer | "
         "new_strategy_thesis_candidate: new family candidate, not yet validated"),
        ("How to read sections 3–6",
         "Section 3 = family/layer-grouped intervention candidates (may be empty, valid). "
         "Section 4 = today's signature candidates (may be 0, valid). "
         "Section 5 = rolling 7-day ledger snapshot, deduped per signature_key. "
         "Section 6 = raw ledger rows for full audit. "
         "Section 7 = raw daily summary for today."),
        ("Canonical identifier",
         "signature_key = canonical grouping/matching key. "
         "signature_candidate_code = human-readable display label. "
         "Do not mix them when building cross-day comparisons."),
    ]
    for rule, explanation in rules:
        _add_row(t, [rule, explanation])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 2 — Multi-Day Executive Summary
# ---------------------------------------------------------------------------

def _u2_multiday_summary(
    doc, normalized_ledger_rows: List[Dict],
    daily_summary: Dict, research_day: str, window_days: int,
) -> None:
    _h2(doc, "2. Multi-Day Executive Summary")

    # Ledger rolling snapshot
    snapshot = build_ledger_snapshot_for_report(normalized_ledger_rows, research_day)
    _p(doc, f"Rolling ledger snapshot ({window_days}-day window ending {research_day}):", bold=True, size=9)

    if not snapshot:
        _p(doc, "No ledger entries within the rolling window.", italic=True, size=9)
    else:
        t = _make_table(doc,
            ["Code", "Support Days", "Recent 7d", "First Seen", "Last Seen", "Status", "Role"],
            [2.2, 0.9, 0.9, 1.0, 1.0, 1.4, 1.5])
        for row in snapshot:
            role = row.get("current_role", "")
            _add_row(t, [
                row.get("signature_candidate_code", row.get("signature_key", "")[:25]),
                row.get("support_days_count", 0),
                row.get("recent_support_days_count", 0),
                row.get("first_seen_date", ""),
                row.get("last_seen_date", ""),
                row.get("latest_validation_status", ""),
                role,
            ], hi_col=6, hi_map=_ROLE_COLORS)

    doc.add_paragraph()

    # Today's summary key fields
    _p(doc, "Today's daily summary (key fields):", bold=True, size=9)
    key_fields = [
        "research_day", "overall_research_health", "research_eligible_cases",
        "dominant_move_class", "dominant_participation_pattern",
        "dominant_structural_quality", "top_candidate_signature_1",
        "old_strategy_improvement_cases_count", "new_strategy_thesis_cases_count",
        "next_research_action", "research_regime", "btc_24h_change_pct", "alt_breadth_pct",
    ]
    t2 = _make_table(doc, ["Field", "Value"], [3.0, 5.4], fill="404040")
    for f in key_fields:
        v = daily_summary.get(f, "")
        if f == "overall_research_health":
            _add_row(t2, [f, str(v) if v is not None else "—"],
                     hi_col=1, hi_map=_HEALTH_COLORS_UA)
        else:
            _add_row(t2, [f, str(v) if v is not None else "—"])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 3 — Daily Intervention Snapshot
# ---------------------------------------------------------------------------

def _u3_intervention_snapshot(
    doc, cases: List[Dict],
    sig_candidates: List[Dict],
    normalized_ledger_rows: List[Dict],
) -> None:
    _h2(doc, "3. Daily Intervention Snapshot")
    _p(doc,
       "Grouped by (strategy_family, issue_layer). Only cases with "
       "decision_grade = OLD_STRATEGY_IMPROVEMENT_CANDIDATE or NEW_STRATEGY_THESIS_CANDIDATE "
       "and research_eligible_YN = Y are included.",
       size=9, italic=True)

    eligible = [c for c in cases if c.get("research_eligible_YN") == "Y"]
    shortlist = build_intervention_shortlist(eligible, sig_candidates, normalized_ledger_rows)

    if not shortlist:
        _p(doc, "No intervention candidates today (no eligible improvement/new-thesis cases).",
           italic=True)
        doc.add_paragraph()
        return

    t = _make_table(doc,
        ["Strategy Family", "Issue Layer", "Cases Today", "Repeated?",
         "Cross-Day Support", "Evidence Source", "Readiness", "Next Action"],
        [1.6, 1.4, 0.8, 0.8, 0.9, 1.3, 1.8, 1.6])
    for row in shortlist:
        _add_row(t, [
            row.get("strategy_family", ""),
            row.get("issue_layer", ""),
            row.get("case_count_today", 0),
            str(row.get("repeated_support_today", 0)),
            row.get("cross_day_support", 0),
            row.get("evidence_source", ""),
            row.get("readiness", ""),
            row.get("recommended_next_action", ""),
        ])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 4 — Daily Signature Candidates
# ---------------------------------------------------------------------------

def _u4_signature_candidates(doc, sig_candidates: List[Dict]) -> None:
    _h2(doc, "4. Daily Signature Candidates")
    _p(doc,
       "Repeated signature candidates: ≥2 eligible cases sharing "
       "(side, pre_move_signature, participation_pattern, structural_quality). "
       "Zero entries = no repeated pattern met threshold today. This is valid and expected "
       "on low-volume or heterogeneous mover days.",
       size=9, italic=True)

    if not sig_candidates:
        _p(doc, "signature_candidates_count = 0  |  reason = no_repeated_pattern_met_threshold")
        doc.add_paragraph()
        return

    _p(doc, f"signature_candidates_count = {len(sig_candidates)}", bold=True)

    headers = ["Code", "N", "Side", "Move Class", "Participation",
               "Struct Quality", "Grade", "Conf", "Validation", "Next Action"]
    col_w   = [2.0, 0.5, 0.5, 1.5, 1.8, 1.5, 2.0, 0.6, 1.2, 1.6]
    t = _make_table(doc, headers, col_w)
    for s in sig_candidates:
        _add_row(t, [
            s.get("signature_candidate_code", ""),
            s.get("support_count", ""),
            s.get("dominant_side", ""),
            s.get("dominant_move_class", ""),
            s.get("dominant_participation_pattern", ""),
            s.get("dominant_structural_quality", ""),
            s.get("decision_grade", ""),
            s.get("confidence", ""),
            s.get("validation_status", ""),
            s.get("next_action", ""),
        ], hi_col=6, hi_map=_GRADE_COLORS)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 5 — Cross-Day Ledger Snapshot
# ---------------------------------------------------------------------------

def _u5_ledger_snapshot(doc, normalized_ledger_rows: List[Dict], research_day: str) -> None:
    _h2(doc, f"5. Cross-Day Ledger Snapshot (as of {research_day})")
    _p(doc,
       "Deduplicated view of signature_evidence_ledger.csv filtered to rolling 7-day window. "
       "One row per signature_key, showing derived support and role.",
       size=9, italic=True)

    snapshot = build_ledger_snapshot_for_report(normalized_ledger_rows, research_day)

    if not snapshot:
        _p(doc, "No ledger data in rolling window.", italic=True)
        doc.add_paragraph()
        return

    t = _make_table(doc,
        ["Code", "First Seen", "Last Seen", "Support Days", "Recent 7d", "Status", "Role"],
        [2.2, 1.0, 1.0, 0.9, 0.8, 1.4, 1.4])
    for row in snapshot:
        role = row.get("current_role", "")
        r = t.add_row()
        vals = [
            row.get("signature_candidate_code", row.get("signature_key", "")[:25]),
            row.get("first_seen_date", ""),
            row.get("last_seen_date", ""),
            row.get("support_days_count", 0),
            row.get("recent_support_days_count", 0),
            row.get("latest_validation_status", ""),
            role,
        ]
        for i, v in enumerate(vals):
            cell = r.cells[i]
            cell.text = str(v) if v is not None else "—"
            _font(cell)
            if i == 6:   # Role column: colour by role
                color = _ROLE_COLORS.get(role)
                if color:
                    _shade(cell, color)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif role == "stale":  # Grey out all cells of stale rows
                _shade(cell, "D9D9D9")
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 6 — Raw Ledger Appendix
# ---------------------------------------------------------------------------

def _u6_raw_ledger(doc, normalized_rows: List[Dict] = None, as_of_day: str = None) -> None:
    _h2(doc, f"6. Raw Ledger Appendix (as of {as_of_day})" if as_of_day else "6. Raw Ledger Appendix")
    _p(doc,
       "All ledger rows as of report day. Key fields only. "
       "Sorted newest first. Same canonical source as Ledger Snapshot (Section 5).",
       size=9, italic=True)
    _p(doc,
       "Status@Write = validation_status recorded at write time, not history-aware. "
       "Use Section 5 (Ledger Snapshot) for normalized current status.",
       size=8, italic=True)

    all_rows_sorted = [r for r in (normalized_rows or [])]

    if not all_rows_sorted:
        _p(doc, "No ledger data available as of this report day.", italic=True)
        doc.add_paragraph()
        return
    all_rows = all_rows_sorted

    key_fields = [
        "research_day", "signature_candidate_code",
        "support_count_day", "confidence_day", "decision_grade_day",
        "validation_status_day", "first_seen_date", "last_seen_date",
    ]
    headers = ["Day", "Code", "Support", "Conf", "Grade", "Status", "First Seen", "Last Seen"]
    col_w   = [1.0, 2.2, 0.8, 0.7, 2.0, 1.3, 1.0, 1.0]
    t = _make_table(doc, headers, col_w, fill="404040")
    for row in sorted(all_rows, key=lambda r: r.get("research_day", ""), reverse=True):
        # Use derived first/last_seen (already normalized by ledger_rows_as_of)
        _add_row(t, [row.get(f, "") for f in key_fields])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 7 — Raw Daily Summary Appendix
# ---------------------------------------------------------------------------

def _u7_raw_summary(doc, daily_summary: Dict) -> None:
    _h2(doc, "7. Raw Daily Summary Appendix")
    _p(doc,
       "All fields from daily_research_summary.csv for today. "
       "Two-column key/value format for easy copy-paste.",
       size=9, italic=True)

    if not daily_summary:
        _p(doc, "No daily summary data provided.", italic=True)
        doc.add_paragraph()
        return

    t = _make_table(doc, ["Field", "Value"], [3.0, 5.4], fill="404040")
    for k, v in daily_summary.items():
        _add_row(t, [k, str(v) if v is not None else "—"])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 8 — Phase 2B Multi-Day Validation Snapshot (compact)
# ---------------------------------------------------------------------------

def _u8_multiday_validation_snapshot(
    doc,
    validation_summary: dict,
    multiday_fam_stats: list,
    family_history_snapshot: Optional[dict] = None,  # accepted but NOT used — Phase 2B independent
) -> None:
    _h2(doc, "8. Phase 2B Multi-Day Validation Snapshot")
    # Phase 2B = today-only conservative promotion frame.
    # Phase 2C family history is in Section 9 and is completely independent.
    # Do NOT override Phase 2B fields with Phase 2C data.
    _p(doc,
       "Phase 2B — today-only conservative promotion frame. "
       "SHORT-side, research_eligible_YN == Y. "
       "Phase 2C family history (multiday) shown in Section 9. "
       "PREPARE_HYPOTHESIS blocked — Phase 2E gates required "
       "(anchor QA, unseen-day validation, sample thresholds). "
       "Bucket-ready: >= 10. Tracking-grade: >= 20. Recommendation: >= 50.",
       size=9, italic=True)

    if not validation_summary:
        _p(doc, "No validation summary available.", size=9, italic=True)
        doc.add_paragraph()
        return

    # --- Promotion state card ---
    promotion_state = validation_summary.get("promotion_state", "DESCRIPTIVE_ONLY")
    t = _make_table(doc, ["Field", "Value"], [3.0, 5.4], fill="404040")

    ps_row = t.add_row()
    ps_row.cells[0].text = "promotion_state"
    ps_row.cells[1].text = promotion_state
    _font(ps_row.cells[0], bold=True)
    _shade(ps_row.cells[1], _PROMOTION_COLORS_UA.get(promotion_state, "808080"))
    for p in ps_row.cells[1].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True

    for label, key in [
        ("Top Family (today only)",                              "top_candidate_family"),
        ("Case Count (today only)",                              "_top_case_count"),
        ("Ledger Research Days (global context only)",           "ledger_research_days_global_context"),
        ("families_at_bucket_ready_or_above (>= 10)",            "families_at_bucket_ready_or_above"),
        ("families_at_tracking_grade_or_above (>= 20)",          "families_at_tracking_grade_or_above"),
        ("families_at_recommendation_grade_or_above (>= 50)",    "families_at_recommendation_grade_or_above"),
        ("promotion_reason",                                     "promotion_reason"),
        ("blocking_reason",                                      "blocking_reason"),
        ("validation_next_step",                                 "validation_next_step"),
    ]:
        _add_row(t, [label, str(validation_summary.get(key, "—"))])
    doc.add_paragraph()

    # --- Per-family breakdown — today-only, no Family Days column (Phase 2C owns that) ---
    if multiday_fam_stats:
        _p(doc, "SHORT family breakdown (today's cases only):", bold=True, size=9)
        tf = _make_table(doc,
            ["Family", "Case Count (today)", "Med F1h", "Med F4h", "Med A4h",
             "Sample Gate", "Confidence", "Sample Note"],
            [2.0, 1.0, 0.7, 0.7, 0.7, 1.1, 0.9, 1.1])
        for row in multiday_fam_stats:
            sg = row.get("sample_gate_status", "—")
            cb = row.get("confidence_band", "—")
            tf_row = tf.add_row()
            vals = [
                row.get("display_family", "—"),
                row.get("case_count", "—"),
                _fmt(row.get("median_f1h")),
                _fmt(row.get("median_f4h")),
                _fmt(row.get("median_a4h")),
                sg, cb, row.get("sample_note", "—"),
            ]
            for i, v in enumerate(vals):
                cell = tf_row.cells[i]
                cell.text = str(v) if v is not None else "—"
                _font(cell)
                if i == 5:  # sample_gate (was 6)
                    color = _SAMPLE_GATE_COLORS_UA.get(str(v))
                    if color:
                        _shade(cell, color)
                        for p in cell.paragraphs:
                            for r in p.runs: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif i == 6:  # confidence_band (was 7)
                    color = _CONFIDENCE_2B_COLORS_UA.get(str(v))
                    if color:
                        _shade(cell, color)
                        for p in cell.paragraphs:
                            for r in p.runs: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph()

    _p(doc,
       "Case Count = today's SHORT eligible cases in this family only. "
       "Phase 2C multiday family history is shown in Section 9. "
       "Ledger Research Days = global context only.",
       size=8, italic=True)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 9 — Phase 2C Family History Foundation
# ---------------------------------------------------------------------------

def _u9_phase2c_foundation(
    doc,
    layer_audit: Optional[dict],
    family_history_snapshot: Optional[dict],
) -> None:
    _h2(doc, "9. Phase 2C Family History Foundation")
    _p(doc,
       "Family history sourced from historical case datasets (daily_case_dataset_{day}.csv). "
       "NOT derived from signature_evidence_ledger. "
       "Layer coverage audit identifies mandatory field gaps from the master spec. "
       "PREPARE_HYPOTHESIS remains blocked — Phase 2E gates required "
       "(anchor QA, unseen-day validation, sample thresholds).",
       size=9, italic=True)

    # --- Layer audit compact summary ---
    if layer_audit:
        _p(doc, "Layer 0–8 Field Coverage:", bold=True, size=9)
        t_lay = _make_table(doc,
            ["Layer", "Name", "Req", "Present", "Missing", "Status"],
            [0.5, 2.5, 0.5, 0.6, 0.6, 0.9])
        _CSTAT_UA = {"OK": "70AD47", "PARTIAL": "FFC000", "BLOCKING": "C00000"}
        for row in layer_audit.get("layer_rows", []):
            status = row.get("coverage_status", "—")
            tr = t_lay.add_row()
            vals = [
                str(row.get("layer_id", "—")),
                row.get("layer_name", "—")[:28],
                str(row.get("required_fields_count", "—")),
                str(row.get("present_fields_count", "—")),
                str(row.get("missing_fields_count", "—")),
                status,
            ]
            for i, v in enumerate(vals):
                cell = tr.cells[i]; cell.text = v; _font(cell)
                if i == 5:
                    color = _CSTAT_UA.get(status)
                    if color:
                        _shade(cell, color)
                        for p in cell.paragraphs:
                            for r in p.runs: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph()
        _p(doc,
           f"Mandatory fields checked: {layer_audit.get('total_required_fields','—')} | "
           f"Present: {layer_audit.get('fields_present_count','—')} | "
           f"Missing: {layer_audit.get('fields_missing_count','—')} | "
           f"Blocking layers: {layer_audit.get('blocking_layers',[])}",
           size=9, italic=True)
        doc.add_paragraph()

    # --- Family history snapshot ---
    if family_history_snapshot:
        _p(doc, "SHORT Family Cross-Day History (Phase 2C):", bold=True, size=9)
        t_fh = _make_table(doc, ["Field", "Value"], [3.0, 5.4], fill="404040")
        for label, key in [
            ("top_candidate_family",               "top_candidate_family"),
            ("top_family_case_count_today",         "top_family_case_count_today"),
            ("top_family_case_count_multiday",      "top_family_case_count_multiday"),
            ("top_family_days_count",               "top_family_days_count"),
            ("family_history_status",               "family_history_status"),
            ("families_with_history_count",         "families_with_history_count"),
            ("History Window Days",                 "history_window_days"),
            ("Historical Case Days Loaded",         "historical_case_days_loaded"),
        ]:
            _add_row(t_fh, [label, str(family_history_snapshot.get(key, "—"))])
        doc.add_paragraph()
        _p(doc, family_history_snapshot.get("history_note", ""), size=8, italic=True)
        doc.add_paragraph()
    else:
        _p(doc, "Family history snapshot not available.", size=9, italic=True)
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2D — Unified pack section helpers
# ---------------------------------------------------------------------------

_BUCKET_COLORS_UA_2D = {
    "bucket_ready":         "70AD47",
    "tracking_grade":       "2E75B6",
    "recommendation_grade": "375623",
    "not_enough_sample":    "808080",
}
_RECOMMEND_COLORS_UA_2D = {
    "keep_tracking":     "2E75B6",
    "descriptive_only":  "808080",
    "blocked":           "C00000",
    "no_change":         "808080",
}


def _u10_phase2d_validation(
    doc,
    family_validation_stats: Optional[list],
) -> None:
    """Section 10 — Phase 2D Family Validation Snapshot."""
    _h2(doc, "10. Phase 2D Family Validation Snapshot")
    _p(doc,
       "Statistical validation per SHORT display_family using multiday case pool. "
       "Bootstrap CI (95%, stdlib). KS/t-test via optional scipy. "
       "Conservative: blocked / not_enough_sample outputs expected at current sample sizes. "
       "PREPARE_HYPOTHESIS not unlocked by Phase 2D stats alone.",
       size=9, italic=True)

    if not family_validation_stats:
        _p(doc,
           "No Phase 2D validation stats available. "
           "Run after historical daily_case_dataset CSVs exist.",
           size=9, italic=True)
        doc.add_paragraph()
        return

    # Compact table: one row per family
    t = _make_table(doc,
        ["Family", "N multi", "Days", "Med F1h", "Med F4h", "Med A4h",
         "CI 95% Low/High", "Bucket", "Stability"],
        [2.0, 0.6, 0.5, 0.65, 0.65, 0.65, 1.4, 0.85, 0.8])

    for row in family_validation_stats:
        bk = row.get("bucket_ready_status", "—")
        st = row.get("stability_flag", "—")
        ci_lo = row.get("bootstrap_ci_low", "na")
        ci_hi = row.get("bootstrap_ci_high", "na")
        is_actionable = row.get("is_actionable_family", True)
        tr = t.add_row()
        vals = [
            row.get("display_family", "—"),
            row.get("case_count_multiday", "—"),
            row.get("family_days_count", "—"),
            _fmt(row.get("median_f1h")),
            _fmt(row.get("median_f4h")),
            _fmt(row.get("median_a4h")),
            f"{ci_lo} / {ci_hi}",
            bk,
            st,
        ]
        for i, v in enumerate(vals):
            cell = tr.cells[i]; cell.text = str(v) if v is not None else "—"; _font(cell)
            if i == 7:
                color = "808080" if not is_actionable else _BUCKET_COLORS_UA_2D.get(str(v))
                if color:
                    _shade(cell, color)
                    for p in cell.paragraphs:
                        for r in p.runs: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # KS/t-test compact table
    t2 = _make_table(doc, ["Family", "KS Status", "t-test Status", "Readiness Note"],
                     [1.8, 1.3, 1.3, 4.0])
    for row in family_validation_stats:
        _add_row(t2, [
            row.get("display_family", "—"),
            str(row.get("ks_test_status_or_na", "—")),
            str(row.get("ttest_status_or_na", "—")),
            str(row.get("statistical_readiness_note", "—"))[:150],
        ])
    doc.add_paragraph()


def _u11_phase2d_contracts(
    doc,
    family_answer_contracts: Optional[list],
) -> None:
    """Section 11 — Phase 2D Family Answer Contracts."""
    _h2(doc, "11. Phase 2D Family Answer Contracts")
    _p(doc,
       "One contract per question family. "
       "Blocked when required Layer fields are missing at source. "
       "Does NOT write live-rule instructions. "
       "See validation_next_step to understand what is needed to unblock.",
       size=9, italic=True)

    if not family_answer_contracts:
        _p(doc,
           "No Phase 2D answer contracts available. "
           "Run after historical case datasets exist.",
           size=9, italic=True)
        doc.add_paragraph()
        return

    t = _make_table(doc,
        ["Question Family", "Rec. State", "Evidence Summary", "Blocking Reason"],
        [1.2, 1.0, 2.5, 3.7])

    for contract in family_answer_contracts:
        rstate = contract.get("recommendation_state", "blocked")
        tr = t.add_row()
        vals = [
            contract.get("family_name", "—"),
            rstate,
            str(contract.get("evidence_summary", "—"))[:130],
            str(contract.get("blocking_reason", "—"))[:130],
        ]
        for i, v in enumerate(vals):
            cell = tr.cells[i]; cell.text = str(v) if v is not None else "—"; _font(cell)
            if i == 1:
                color = _RECOMMEND_COLORS_UA_2D.get(str(v))
                if color:
                    _shade(cell, color)
                    for p in cell.paragraphs:
                        for r in p.runs: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()

    # Validation next step per contract
    _p(doc, "Validation next steps per question family:", bold=True, size=9)
    t3 = _make_table(doc, ["Question Family", "Validation Next Step"], [1.5, 6.9])
    for contract in family_answer_contracts:
        _add_row(t3, [
            contract.get("family_name", "—"),
            str(contract.get("validation_next_step", "—")),
        ])
    doc.add_paragraph()




_P2E_STATE_COLORS_UA = {
    "READY_FOR_CONTROLLED_VALIDATION":     "375623",
    "PREPARE_HYPOTHESIS":                  "538135",
    "NOT_READY_FOR_CONTROLLED_VALIDATION": "C55A11",
    "KEEP_TRACKING":                       "2E75B6",
    "DESCRIPTIVE_ONLY":                    "808080",
}

_P2E_UNSEEN_COLORS_UA = {
    "NOT_READY_no_holdout_data":             "808080",
    "NOT_READY_insufficient_holdout_sample": "C55A11",
    "PARTIAL_holdout_available":             "2E75B6",
}


def _trunc_p2e(text, limit: int = 130) -> str:
    """Clean truncation with ellipsis for Phase 2E display fields."""
    s = str(text) if text is not None else "—"
    if len(s) <= limit:
        return s
    cut = s[:limit - 1].rsplit(" ", 1)[0]
    return cut + "…"


def _u12_phase2e_controlled_validation(
    doc,
    controlled_validation_state: Optional[list],
    unseen_day_summary: Optional[list],
) -> None:
    """Section 12 — Phase 2E Controlled Validation Gate.
    Renderer only: consumes producer outputs, does not invent gate logic.
    """
    _h2(doc, "12. Phase 2E — Controlled Validation Gate")
    _p(doc,
       "Gate-building phase. Expected output: conservative for all families. "
       "PREPARE_HYPOTHESIS / READY_FOR_CONTROLLED_VALIDATION require all hard gates to pass. "
       "Conservative result is correct.",
       size=9, italic=True)

    if not controlled_validation_state:
        _p(doc,
           "No Phase 2E controlled validation state available. "
           "Requires Phase 2C family history + Phase 2D validation stats.",
           size=9, italic=True)
        doc.add_paragraph()
    else:
        _p(doc, "Controlled Validation State per display_family:", bold=True, size=9)
        t = _make_table(
            doc,
            ["Family", "State", "N multi", "Days", "Anchor Gate", "Blocker Summary"],
            [1.8, 2.0, 0.6, 0.5, 1.2, 2.3],
        )
        for row in controlled_validation_state:
            state = row.get("controlled_validation_state", "—")
            tr = t.add_row()
            vals = [
                row.get("display_family", "—"),
                state,
                row.get("case_count_multiday", "—"),
                row.get("family_days_count", "—"),
                row.get("anchor_qa_gate", "—"),
                _trunc_p2e(row.get("promotion_blocker_summary", "—"), 130),
            ]
            for i, v in enumerate(vals):
                cell = tr.cells[i]
                cell.text = str(v) if v is not None else "—"
                _font(cell, size=8)
                if i == 1:
                    color = _P2E_STATE_COLORS_UA.get(str(v))
                    if color:
                        _shade(cell, color)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph()

    if not unseen_day_summary:
        _p(doc, "No unseen-day validation summary available.", size=9, italic=True)
        doc.add_paragraph()
        return

    _p(doc, "Unseen-Day / Holdout Validation Summary:", bold=True, size=9)
    t2 = _make_table(
        doc,
        ["Family", "Unseen Status", "Holdout N", "Days", "Summary"],
        [1.8, 2.0, 0.7, 0.5, 3.4],
    )
    for row in unseen_day_summary:
        status = row.get("unseen_validation_status", "—")
        tr = t2.add_row()
        vals = [
            row.get("display_family", "—"),
            status,
            row.get("holdout_case_count", 0),
            row.get("holdout_days_used", 0),
            _trunc_p2e(row.get("holdout_result_summary", "—"), 130),
        ]
        for i, v in enumerate(vals):
            cell = tr.cells[i]
            cell.text = str(v) if v is not None else "—"
            _font(cell, size=8)
            if i == 1:
                color = _P2E_UNSEEN_COLORS_UA.get(str(v))
                if color:
                    _shade(cell, color)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()
    _p(doc,
       "> Unseen-day validation not yet run. "
       "Holdout data required before any promotion claim is valid.",
       size=8, italic=True)
    doc.add_paragraph()

# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def build_unified_analysis_pack(
    research_day: str,
    cases: List[Dict],
    sig_candidates: List[Dict],
    daily_summary: Dict,
    output_path: str,
    window_days: int = 7,
    layer_audit: Optional[dict] = None,
    family_history_snapshot: Optional[dict] = None,
    p2d_family_validation_stats: Optional[list] = None,
    p2d_family_answer_contracts: Optional[list] = None,
    p2e_controlled_validation_state: Optional[list] = None,
    p2e_unseen_day_summary: Optional[list] = None,
) -> str:
    """
    Build R1_unified_analysis_pack_<research_day>.docx

    Sections 1–12. Sections 10–11 are Phase 2D. Section 12 is Phase 2E gate.
    Downstream-only. Does not touch live runtime files.
    Safe on first run (ledger may not exist yet).

    Returns the output_path on success.
    """
    # Canonical as-of-day ledger source — used by ALL ledger-derived sections
    normalized_ledger_rows = ledger_rows_as_of(research_day, window_days=window_days)

    # Phase 2B derivations — consume helpers from signature_ledger, no re-derive here
    _ledger_days_context = len(set(
        r.get("research_day", "") for r in normalized_ledger_rows
        if r.get("research_day")
    ))
    multiday_fam_stats   = build_multiday_family_stats(cases, normalized_ledger_rows)
    multiday_inter_stats = build_multiday_interaction_stats(cases)
    validation_summary   = build_controlled_validation_summary(
        multiday_fam_stats, multiday_inter_stats, _ledger_days_context,
    )
    if multiday_fam_stats:
        _top_fam_u = validation_summary.get("top_candidate_family", "")
        _top_row_u = next((r for r in multiday_fam_stats if r.get("display_family") == _top_fam_u), None)
        validation_summary["_top_case_count"] = (
            _top_row_u.get("case_count", "—") if _top_row_u else "—"
        )

    doc = Document()

    # Cover
    _h1(doc, "R1 Unified Analysis Pack")
    _p(doc, f"Research Day: {research_day}", bold=True)
    _p(doc, f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    _p(doc,
       f"Cases: {len(cases)}  |  "
       f"Sig Candidates: {len(sig_candidates)}  |  "
       f"Ledger Rows: {len(normalized_ledger_rows)}  |  "
       f"Window: {window_days} days")
    _p(doc,
       "This file is downstream-only. It does not modify live bot config, "
       "lifecycle, strategy, or runtime files.",
       size=9, italic=True)
    doc.add_page_break()

    _u1_interpretation_guide(doc)
    doc.add_page_break()

    _u2_multiday_summary(doc, normalized_ledger_rows, daily_summary, research_day, window_days)
    _u3_intervention_snapshot(doc, cases, sig_candidates, normalized_ledger_rows)
    doc.add_page_break()

    _u4_signature_candidates(doc, sig_candidates)
    _u5_ledger_snapshot(doc, normalized_ledger_rows, research_day)
    doc.add_page_break()

    _u6_raw_ledger(doc, normalized_rows=normalized_ledger_rows, as_of_day=research_day)
    doc.add_page_break()

    _u7_raw_summary(doc, daily_summary)

    _u8_multiday_validation_snapshot(
        doc, validation_summary, multiday_fam_stats, family_history_snapshot)

    _u9_phase2c_foundation(doc, layer_audit, family_history_snapshot)

    _u10_phase2d_validation(doc, p2d_family_validation_stats)
    _u11_phase2d_contracts(doc, p2d_family_answer_contracts)
    _u12_phase2e_controlled_validation(
        doc, p2e_controlled_validation_state, p2e_unseen_day_summary)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

"""
research/top_movers/docx_report_builder.py

Builds the full daily research pack DOCX using python-docx.

Fix: Sections 3 and 10 now explicitly distinguish:
  A. Repeated signature candidates (require >= threshold repeated cases → may be zero)
  B. Case-level strategy theses (per-case interpretation → may exist even when A = 0)
"""

import os
from datetime import datetime, timezone
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from research.top_movers.io import image_path
from research.top_movers.signature_ledger import (
    load_and_normalize_ledger_rows,
    build_ledger_snapshot_for_report,
    validate_ledger_semantics,
    build_intervention_shortlist,
    build_measurement_decision_card,
    build_trusted_weak_deferred,
    # Phase 2A
    build_anchor_qa_summary,
    build_short_family_boards,
    build_interaction_board,
    build_measurement_confidence_summary,
    # Phase 2B
    build_multiday_family_stats,
    build_multiday_interaction_stats,
    build_controlled_validation_summary,
    # Phase 2D
    build_family_validation_stats,
    build_family_answer_contracts,
)

# Maps image key to primary anchor code (for missing-reason lookup)
_IMG_KEY_TO_ANCHOR = {
    "P0_context_1h":            "P0",
    "P0_P1_setup_15m":          "P0",
    "P1_ignition_5m":           "P1",
    "P2_P3_break_expansion_5m": "P2",
    "P4_resolution_15m":        "P4",
}

_IMAGE_DEFS = [
    ("P0_context_1h",              "P0 Context (1h)"),
    ("P0_P1_setup_15m",            "P0→P1 Setup (15m)"),
    ("P1_ignition_5m",             "P1 Ignition (5m)"),
    ("P2_P3_break_expansion_5m",   "P2→P3 Break+Expansion (5m)"),
    ("P4_resolution_15m",          "P4 Resolution (15m)"),
]

_GRADE_COLORS = {
    "OLD_STRATEGY_IMPROVEMENT_CANDIDATE": "2E75B6",
    "NEW_STRATEGY_THESIS_CANDIDATE":      "70AD47",
    "KEEP_TRACKING":                      "FFC000",
    "DESCRIPTIVE_ONLY":                   "808080",
    "NOT_RELIABLE_YET":                   "C00000",
}
_HEALTH_COLORS = {"CLEAN": "70AD47", "CLEAN_WITH_VISUAL_GAPS": "92D050", "PARTIAL": "FFC000", "WEAK": "C00000"}
_YN_COLORS     = {"Y": "70AD47", "N": "C00000"}
_DQ_COLORS     = {"CLEAN": "70AD47", "PARTIAL": "FFC000", "WEAK": "C00000"}

# Phase 2A color maps
_READINESS_COLORS = {
    "MEASUREMENT_READY":   "70AD47",
    "USABLE_WITH_CAUTION": "FFC000",
    "DIRECTIONAL_ONLY":    "2E75B6",
    "NOT_RUN":             "C00000",
}
_CONFIDENCE_BAND_COLORS = {
    "HIGH":   "70AD47",
    "MEDIUM": "FFC000",
    "LOW":    "C00000",
}
_GATE_COLORS = {
    "MET":     "70AD47",
    "PARTIAL": "FFC000",
    "NOT_MET": "C00000",
}
_STABILITY_COLORS_P2 = {
    "STABLE":       "70AD47",
    "EARLY_SIGNAL": "FFC000",
    "UNSTABLE":     "C00000",
}
_REGIME_COLORS_P2 = {
    "CONSISTENT": "70AD47",
    "MIXED":      "FFC000",
    "UNKNOWN":    "808080",
}

# Phase 2B color maps
_SAMPLE_GATE_COLORS_2B = {
    "RECOMMENDATION_GRADE": "70AD47",
    "TRACKING_GRADE":       "FFC000",
    "LOW_SAMPLE":           "C0504D",
    "NOT_ENOUGH_SAMPLE":    "808080",
}
_BUCKET_GATE_COLORS_2B = {
    "BUCKET_READY":         "70AD47",
    "BUCKET_THIN":          "FFC000",
    "BUCKET_INSUFFICIENT":  "808080",
}
_CONFIDENCE_2B_COLORS = {
    "MODERATE":         "FFC000",
    "LOW":              "C00000",
    "DESCRIPTIVE_ONLY": "808080",
}
_PROMOTION_COLORS = {
    "DESCRIPTIVE_ONLY":                    "808080",
    "KEEP_TRACKING":                       "FFC000",
    "PREPARE_HYPOTHESIS":                  "2E75B6",
    "NOT_READY_FOR_CONTROLLED_VALIDATION": "C00000",
}


def _shade(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _white_bold(cell):
    for p in cell.paragraphs:
        for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

def _font(cell, bold=False, size=8):
    for p in cell.paragraphs:
        for r in p.runs: r.bold=bold; r.font.size=Pt(size)

def _make_table(doc, headers, col_w, fill="2E75B6"):
    if sum(col_w) > 6.5: col_w = [w * 6.5 / sum(col_w) for w in col_w]
    t = doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    tbl = t._tbl; tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"),str(int(sum(col_w)*1440))); tblW.set(qn("w:type"),"dxa"); tblPr.append(tblW)
    for i, cell in enumerate(t.rows[0].cells):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        tcW=OxmlElement("w:tcW"); tcW.set(qn("w:w"),str(int(col_w[i]*1440))); tcW.set(qn("w:type"),"dxa"); tcPr.insert(0,tcW)
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.text=h; _shade(cell,fill); _white_bold(cell)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    return t

def _add_row(t, vals, hi_col=-1, hi_map=None):
    row=t.add_row()
    for i,v in enumerate(vals):
        cell=row.cells[i]; cell.text=str(v) if v is not None else "—"; _font(cell)
        if hi_map and i==hi_col:
            color=hi_map.get(str(v))
            if color:
                _shade(cell,color)
                for p in cell.paragraphs:
                    for r in p.runs: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

def _h1(doc, t): doc.add_heading(t, level=1)
def _h2(doc, t): doc.add_heading(t, level=2)
def _p(doc, t, bold=False, size=10, italic=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.font.size=Pt(size); r.italic=italic; return p

def _dom(cases, f):
    vs=[c.get(f,"") for c in cases if c.get(f)]; return max(set(vs),key=vs.count) if vs else "—"
def _pct(n,total): return f"{round(100*n/max(total,1))}%"
def _fmt(v, nd=3):
    if v is None or v=="" or v=="None": return "—"
    try: return str(round(float(v),nd))
    except: return str(v)


# ---------------------------------------------------------------------------
# Section 1 — Data Quality Gate
# ---------------------------------------------------------------------------

def _s1_data_quality(doc, cases):
    _h2(doc, "1. Data Quality Gate")
    eligible  = sum(1 for c in cases if c.get("research_eligible_YN")=="Y")
    full_vis  = sum(1 for c in cases if c.get("full_visual_complete_YN")=="Y")
    proxy_ok  = sum(1 for c in cases if c.get("proxy_complete_YN")=="Y")
    outcome_ok= sum(1 for c in cases if c.get("outcome_complete_YN")=="Y")
    _p(doc, f"Research eligible: {eligible}/{len(cases)}  |  Full visual: {full_vis}/{len(cases)}  |  Proxy ≥50%: {proxy_ok}/{len(cases)}  |  Outcomes: {outcome_ok}/{len(cases)}")
    t=_make_table(doc, ["Symbol","Side","Res.Elig","FullVis","Proxy","Outcome","DQ Flag","Note"],
                  [0.9,0.5,0.8,0.7,0.6,0.7,0.8,2.7])
    for c in cases:
        _add_row(t, [c.get("symbol"),c.get("side"),c.get("research_eligible_YN"),
            c.get("full_visual_complete_YN"),c.get("proxy_complete_YN"),c.get("outcome_complete_YN"),
            c.get("data_quality_flag"),c.get("data_quality_note") or ""],
            hi_col=2, hi_map=_YN_COLORS)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 2 — Research Integrity Panel
# ---------------------------------------------------------------------------

def _s2_integrity(doc, cases, image_results_all):
    _h2(doc, "2. Research Integrity Panel")
    eligible    = [c for c in cases if c.get("research_eligible_YN")=="Y"]
    missing_img = [c.get("symbol","") for c in cases if c.get("full_visual_complete_YN")!="Y"]
    missing_prx = [c.get("symbol","") for c in cases if c.get("proxy_complete_YN")!="Y"]
    missing_out = [c.get("symbol","") for c in cases if c.get("outcome_complete_YN")!="Y"]
    n=len(cases); er=len(eligible)/n if n else 0; pr=(n-len(missing_prx))/n if n else 0; ir=(n-len(missing_img))/n if n else 0
    health = ("CLEAN_WITH_VISUAL_GAPS" if missing_img else "CLEAN") if er>=0.75 and pr>=0.60 else "PARTIAL" if er>=0.40 or pr>=0.40 else "WEAK"
    t=_make_table(doc, ["Metric","Value"], [3.0,4.7], fill="404040")
    _add_row(t, ["Research health", health], hi_col=1, hi_map=_HEALTH_COLORS)
    _add_row(t, ["Research-eligible", f"{len(eligible)}/{n}"])
    _add_row(t, ["Missing any image (full_visual_complete=N)", ", ".join(missing_img) or "none"])
    _add_row(t, ["Missing proxy history (proxy_complete=N)",   ", ".join(missing_prx) or "none"])
    _add_row(t, ["Missing outcome horizon (outcome_complete=N)",", ".join(missing_out) or "none"])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 3 — Executive Research Summary
# Note: distinguish repeated sig candidates vs case-level theses (fix 3)
# ---------------------------------------------------------------------------

def _s3_exec_summary(doc, cases, selection_context, signature_candidates):
    _h2(doc, "3. Executive Research Summary")
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    if not eligible: _p(doc,"No eligible cases today."); return

    n=len(cases); missing_prx=sum(1 for c in cases if c.get("proxy_complete_YN")!="Y")
    missing_img=sum(1 for c in cases if c.get("full_visual_complete_YN")!="Y")
    er=len(eligible)/n if n else 0; pr=(n-missing_prx)/n if n else 0; ir=(n-missing_img)/n if n else 0
    health = ("CLEAN_WITH_VISUAL_GAPS" if missing_img>0 else "CLEAN") if er>=0.75 and pr>=0.60 else "PARTIAL" if er>=0.40 or pr>=0.40 else "WEAK"

    improve=sum(1 for c in eligible if c.get("strategy_action_type")=="improve_existing")
    create =sum(1 for c in eligible if c.get("strategy_action_type")=="create_new")
    observe=sum(1 for c in eligible if c.get("strategy_action_type")=="keep_observing")

    n_sigs = len(signature_candidates)
    top_sig = signature_candidates[0].get("signature_candidate_code","—") if signature_candidates else "none today"

    t=_make_table(doc, ["Key","Value"], [3.0,4.7])
    rows=[
        ("Research health", health),
        ("Regime", selection_context.get("research_regime","—")),
        ("BTC 24h", f"{_fmt(selection_context.get('btc_24h_change_pct'),2)}%"),
        ("Alt Breadth", f"{_fmt(selection_context.get('alt_breadth_pct'),1)}%"),
        ("Dominant move class", _dom(eligible,"move_class")),
        ("Dominant pre-move signature", _dom(eligible,"pre_move_signature")),
        ("Dominant participation", _dom(eligible,"participation_pattern")),
        ("Dominant structural quality", _dom(eligible,"structural_quality")),
        ("improve_existing (case-level)", str(improve)),
        ("create_new / new thesis (case-level)", str(create)),
        ("keep_observing (case-level)", str(observe)),
        (f"Repeated sig candidates (≥{2} cases)", str(n_sigs)),
        ("Top repeated signature", top_sig),
    ]
    for k,v in rows:
        _add_row(t, [k,v], hi_col=1, hi_map={"CLEAN":"70AD47","PARTIAL":"FFC000","WEAK":"C00000"})
    doc.add_paragraph()

    # Distinction note (fix 3A)
    _p(doc,
       "DISTINCTION — Repeated Signature Candidates vs Case-Level Strategy Theses:",
       bold=True, size=9)
    _p(doc,
       "  A. Repeated signature candidates (Section 10): require >= 2 eligible cases to share the "
       "same (side, pre_move_signature, participation_pattern, structural_quality). "
       "Zero candidates today means no pattern repeated — this is valid and expected on low-volume days.",
       size=9, italic=True)
    _p(doc,
       "  B. Case-level strategy theses (Section 13): derived from individual case interpretation. "
       "A case may have decision_grade = NEW_STRATEGY_THESIS_CANDIDATE or OLD_STRATEGY_IMPROVEMENT_CANDIDATE "
       "even when zero repeated signature candidates exist. These are per-case conclusions, not confirmed repeated patterns.",
       size=9, italic=True)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Sections 4-9
# ---------------------------------------------------------------------------

def _s4_selection_board(doc, cases):
    _h2(doc, "4. Selection Board")
    t=_make_table(doc, ["#","Symbol","Side","Return%","Rank","Regime","Eligible","FullVis"],
                  [0.4,1.0,0.5,0.8,0.5,1.5,0.7,0.7])
    for i,c in enumerate(cases,1):
        _add_row(t, [i,c.get("symbol"),c.get("side"),_fmt(c.get("daily_return_pct"),2),
            c.get("top_mover_rank"),c.get("research_regime",""),
            c.get("research_eligible_YN"),c.get("full_visual_complete_YN")],
            hi_col=6, hi_map=_YN_COLORS)
    doc.add_paragraph()


def _s5_move_archetype(doc, cases):
    _h2(doc, "5. Move Archetype Distribution")
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    counts: Dict[str,int]={}
    for c in eligible: mc=c.get("move_class","—"); counts[mc]=counts.get(mc,0)+1
    t=_make_table(doc, ["Move Class","Count","Share"], [3.5,0.8,0.8])
    for mc,cnt in sorted(counts.items(), key=lambda x:-x[1]):
        _add_row(t, [mc, cnt, _pct(cnt,len(eligible))])
    doc.add_paragraph()


def _s6_pre_move_sig(doc, cases):
    _h2(doc, "6. Pre-Move Signature Board")
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    groups: Dict[str,list]={}
    for c in eligible: groups.setdefault(c.get("pre_move_signature","—"),[]).append(c)
    t=_make_table(doc, ["Signature","Count","Share","Dominant Side"], [2.5,0.7,0.7,1.5])
    for sig,cs in sorted(groups.items(), key=lambda x:-len(x[1])):
        _add_row(t, [sig,len(cs),_pct(len(cs),len(eligible)),_dom(cs,"side")])
    doc.add_paragraph()


def _s7_participation(doc, cases):
    _h2(doc, "7. Participation Pattern Board")
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    counts: Dict[str,int]={}
    for c in eligible: counts[c.get("participation_pattern","—")]=counts.get(c.get("participation_pattern","—"),0)+1
    t=_make_table(doc, ["Pattern","Count","Share"], [3.5,0.8,0.8])
    for pat,cnt in sorted(counts.items(), key=lambda x:-x[1]):
        _add_row(t, [pat,cnt,_pct(cnt,len(eligible))])
    doc.add_paragraph()


def _s8_structural_quality(doc, cases):
    _h2(doc, "8. Structural Quality Board")
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    counts: Dict[str,int]={}
    for c in eligible: counts[c.get("structural_quality","—")]=counts.get(c.get("structural_quality","—"),0)+1
    t=_make_table(doc, ["Structural Quality","Count","Share"], [3.5,0.8,0.8])
    for sq,cnt in sorted(counts.items(), key=lambda x:-x[1]):
        _add_row(t, [sq,cnt,_pct(cnt,len(eligible))])
    doc.add_paragraph()


def _s9_outcome_quality(doc, cases):
    _h2(doc, "9. Outcome Quality Board")
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    groups: Dict[str,list]={}
    for c in eligible: groups.setdefault(c.get("resolution_label","—"),[]).append(c)
    t=_make_table(doc, ["Resolution Label","Count","Share","Avg f1h","Avg f4h","Avg a4h"],
                  [2.5,0.6,0.6,0.9,0.9,0.9])
    for rl,cs in sorted(groups.items(), key=lambda x:-len(x[1])):
        _add_row(t, [rl,len(cs),_pct(len(cs),len(eligible)),
            _fmt(_avg_field(cs,"future_1h_max_favor_pct")),
            _fmt(_avg_field(cs,"future_4h_max_favor_pct")),
            _fmt(_avg_field(cs,"future_4h_max_adverse_pct"))])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 10 — Signature Evidence Board
# Note: explicitly states distinction and handles zero-candidate case (fix 3B)
# ---------------------------------------------------------------------------

def _s10_signature_evidence(doc, signature_candidates):
    _h2(doc, "10. Signature Evidence Board")

    # Distinction note (fix 3B)
    _p(doc,
       "Repeated signature candidates: require >= 2 eligible cases with the same "
       "(side, pre_move_signature, participation_pattern, structural_quality). "
       "Zero entries is valid — it means no cross-case pattern repeated today.",
       size=9, italic=True)
    _p(doc,
       "Individual cases may still have decision_grade = NEW_STRATEGY_THESIS_CANDIDATE or "
       "OLD_STRATEGY_IMPROVEMENT_CANDIDATE (see Section 13). Those are case-level theses, "
       "not confirmed repeated signatures.",
       size=9, italic=True)

    if not signature_candidates:
        _p(doc, f"No repeated signature candidates today (threshold: ≥2 cases).")
        _p(doc, "signature_candidates_count = 0  |  reason = no_repeated_pattern_met_threshold", size=9)
        doc.add_paragraph(); return

    _p(doc, f"signature_candidates_count = {len(signature_candidates)}")
    t=_make_table(doc,
        ["Code","Description","N","Share%","Side","MClass","PartPat","Struct","Med1hF","Med4hF","Action","Conf","NextAct"],
        [1.3,2.5,0.4,0.6,0.5,1.5,1.8,1.5,0.7,0.7,1.2,0.6,1.5])
    for s in signature_candidates:
        _add_row(t, [
            s.get("signature_candidate_code"), s.get("signature_description"),
            s.get("support_count"), s.get("support_share_pct"),
            s.get("dominant_side"), s.get("dominant_move_class"),
            s.get("dominant_participation_pattern"), s.get("dominant_structural_quality"),
            _fmt(s.get("median_1h_favor")), _fmt(s.get("median_4h_favor")),
            s.get("maps_to_existing_strategy_family"), s.get("confidence"), s.get("next_action"),
        ])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Sections 11-18
# ---------------------------------------------------------------------------

def _s11_strategy_mapping(doc, cases):
    _h2(doc, "11. Strategy Mapping Board")
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    t=_make_table(doc, ["Symbol","Side","Move Class","Maps To","Fit","Layer","Hypothesis","Cand?"],
                  [0.9,0.5,1.8,1.5,0.5,1.3,2.0,0.6])
    for c in eligible:
        _add_row(t, [c.get("symbol"),c.get("side"),c.get("move_class"),
            c.get("maps_to_existing_strategy_family"),c.get("existing_strategy_fit_confidence"),
            c.get("improvement_target_layer"),c.get("improvement_hypothesis"),
            c.get("existing_strategy_improvement_candidate_YN")])
    doc.add_paragraph()


def _s12_new_strategy(doc, cases):
    _h2(doc, "12. Research Families Under Investigation")
    eligible=[c for c in cases if c.get("new_strategy_candidate_flag")=="Y"]
    _p(doc,
       "These are exploratory research families under investigation — NOT actual NEW_STRATEGY_THESIS_CANDIDATE cases. "
       "A case reaching this board means its move pattern did not map to an existing strategy and showed "
       "some structural interest. It does NOT mean a new strategy is validated or recommended.",
       size=9, italic=True)
    if not eligible: _p(doc,"No exploratory families today."); doc.add_paragraph(); return
    t=_make_table(doc, ["Symbol","Side","Candidate Family","Trigger","Env","Invalid","Conf"],
                  [0.9,0.5,1.8,2.0,1.3,1.5,0.7])
    for c in eligible:
        _add_row(t, [c.get("symbol"),c.get("side"),c.get("candidate_strategy_family_name"),
            c.get("candidate_trigger_description"),c.get("candidate_environment"),
            c.get("candidate_invalid_pattern"),c.get("new_strategy_confidence")])
    doc.add_paragraph()


def _s13_decision_grade(doc, cases):
    _h2(doc, "13. Decision Grade Board")
    _p(doc, "Decision grades are per-case (case-level theses). These exist independently of "
            "repeated signature candidates in Section 10.", size=9, italic=True)
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    t=_make_table(doc, ["Symbol","Side","Decision Grade","Action","Conf","Prx%","Out%","Vis","Caution"],
                  [0.9,0.5,2.2,1.4,0.7,0.6,0.6,0.6,0.7])
    for c in eligible:
        _add_row(t, [c.get("symbol"),c.get("side"),c.get("decision_grade"),
            c.get("strategy_action_type"),c.get("classification_confidence"),
            _fmt(c.get("proxy_completeness_score"),0),
            _fmt(c.get("outcome_completeness_score"),0),
            c.get("full_visual_complete_YN"),c.get("caution_flag")],
            hi_col=2, hi_map=_GRADE_COLORS)
    doc.add_paragraph()


def _s14_trap_caution(doc, cases):
    _h2(doc, "14. Trap / Caution Board")
    traps=_detect_traps(cases)
    if not traps: _p(doc,"No trap patterns today."); doc.add_paragraph(); return
    t=_make_table(doc, ["Trap Code","Description","Affected","Why Risky","Caution"], [1.5,2.0,1.5,2.0,1.7])
    for tr in traps:
        _add_row(t, [tr["trap_code"],tr["description"],tr["affected"],tr["why_risky"],tr["caution"]])
    doc.add_paragraph()


def _s15_review_queue(doc, cases):
    _h2(doc, "15. Human Review Queue")
    queue=[c for c in cases if c.get("decision_grade") in
           ("OLD_STRATEGY_IMPROVEMENT_CANDIDATE","NEW_STRATEGY_THESIS_CANDIDATE")
           or c.get("caution_flag")=="Y"]
    if not queue: _p(doc,"No cases require immediate review."); doc.add_paragraph(); return
    t=_make_table(doc, ["Case ID","Symbol","Side","Decision Grade","Reason"], [2.0,1.0,0.5,2.5,2.7])
    for c in queue:
        reason=[]
        if c.get("decision_grade") in ("OLD_STRATEGY_IMPROVEMENT_CANDIDATE","NEW_STRATEGY_THESIS_CANDIDATE"):
            reason.append(c.get("decision_grade",""))
        if c.get("caution_flag")=="Y": reason.append("caution_flag")
        _add_row(t, [c.get("case_id"),c.get("symbol"),c.get("side"),c.get("decision_grade")," | ".join(reason)])
    doc.add_paragraph()


def _s16_case_registry(doc, cases):
    _h2(doc, "16. Case Registry Preview")
    t=_make_table(doc, ["Case ID","Symbol","Side","Return%","Move Class","Struct","Flow Phase","Resolution","Grade"],
                  [1.8,0.9,0.5,0.7,1.5,1.5,1.8,1.8,2.2])
    for c in cases:
        _add_row(t, [c.get("case_id"),c.get("symbol"),c.get("side"),_fmt(c.get("daily_return_pct"),2),
            c.get("move_class"),c.get("structural_quality"),c.get("flow_phase_code"),
            c.get("resolution_label"),c.get("decision_grade")],
            hi_col=8, hi_map=_GRADE_COLORS)
    doc.add_paragraph()


def _s17_case_appendix(doc, cases, anchor_rows, research_day, image_results_all=None):
    _h1(doc, "17. Full Case Detail Appendix")
    _p(doc, "One section per token. 5 images embedded. v2 taxonomy labels.", size=9)

    anchor_by_case: Dict[str,List[Dict]]={}
    for row in anchor_rows:
        anchor_by_case.setdefault(row.get("case_id",""),[]).append(row)

    for c in cases:
        doc.add_page_break()
        cid=c.get("case_id",""); sym=c.get("symbol",""); side=c.get("side","")
        _h2(doc, f"{sym} {side} | {cid}")
        _p(doc, f"Return: {_fmt(c.get('daily_return_pct'),2)}%  |  Move: {c.get('move_class')}  |  "
                f"Struct: {c.get('structural_quality')}  |  Grade: {c.get('decision_grade')}", bold=True)
        _p(doc, f"Elig={c.get('research_eligible_YN')}  FullVis={c.get('full_visual_complete_YN')}  "
                f"Proxy={c.get('proxy_complete_YN')}  Outcome={c.get('outcome_complete_YN')}  "
                f"Caution={c.get('caution_flag')}", size=9)
        if c.get("full_visual_complete_YN") != "Y":
            _p(doc,
               "Note: full_visual_complete = N means one or more chart images could not be rendered. "
               "This does NOT affect outcome/resolution fields (which are computed from price data, not images). "
               "P4 anchor availability and P4 image availability are separate — see anchor table above.",
               size=8, italic=True)
        t1=c.get("case_takeaway_1",""); t2=c.get("case_takeaway_2","")
        if t1: _p(doc, f"► {t1}")
        if t2: _p(doc, f"  {t2}", size=9)

        _p(doc, "Anchor Timeline:", bold=True, size=9)
        at=_make_table(doc, ["Anchor","Ts ms","Close","Quality","Image"], [0.8,1.8,1.2,1.8,1.0], fill="404040")
        for code in ["P0","P1","P2","P3","P4"]:
            a=next((r for r in anchor_by_case.get(cid,[]) if r.get("anchor_code")==code), None)
            if a:
                q=""
                if code=="P0": q=f"comp={_fmt(a.get('compression_score'))}"
                elif code=="P2": q=f"bq={a.get('break_quality_band','')} ({_fmt(a.get('break_quality_score'))})"
                elif code=="P3": q=f"ext={_fmt(a.get('directional_extension_pct'))}%"
                _add_row(at, [code,a.get("anchor_ts_ms"),a.get("bar_close"),q,a.get("image_created_YN","N")])
        doc.add_paragraph()

        _p(doc, "Proxy & Outcome:", bold=True, size=9)
        pt=_make_table(doc, ["Field","Value","Field","Value"], [2.0,1.5,2.0,1.5], fill="404040")
        pairs=[
            ("flow_phase",c.get("flow_phase_code")), ("resolution",c.get("resolution_label")),
            ("participation",c.get("participation_pattern")), ("pre_sig",c.get("pre_move_signature")),
            ("large_proxy",_fmt(c.get("large_participant_proxy"))), ("crowd_proxy",_fmt(c.get("crowd_participation_proxy"))),
            ("f1h_favor",_fmt(c.get("future_1h_max_favor_pct"))), ("f4h_favor",_fmt(c.get("future_4h_max_favor_pct"))),
        ]
        for i in range(0,len(pairs),2):
            k1,v1=pairs[i]; k2,v2=pairs[i+1] if i+1<len(pairs) else ("","")
            _add_row(pt, [k1,v1,k2,v2])
        doc.add_paragraph()

        _p(doc, "Charts:", bold=True, size=9)
        for img_key, img_label in _IMAGE_DEFS:
            fp=image_path(research_day, cid, img_key)
            _p(doc, f"  {img_label}", size=8)
            if os.path.exists(fp):
                try:
                    doc.add_picture(fp, width=Inches(6.5))
                    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    _p(doc, f"  [Embed failed: {e}]", size=8)
            else:
                _img_res = (image_results_all or {}).get(cid, {}).get(img_key, {})
                _reason = _img_res.get("reason", "") or "reason_unknown"
                _p(doc, f"  [Missing: {img_key} | reason: {_reason}]", size=8)
            doc.add_paragraph()


def _s18_footer(doc, cases, signature_candidates):
    doc.add_page_break(); _h1(doc, "Research Action Footer")
    eligible=[c for c in cases if c.get("research_eligible_YN")=="Y"]
    # All repeated signature candidates today are 'keep tracking' by definition —
    # filtering by strategy_action_type=='keep_observing' uses the wrong field
    # (that is a case-level field, not a signature-level action value).
    keep  = [s.get("signature_candidate_code","") for s in signature_candidates]
    imp   = [c.get("case_id","") for c in eligible if c.get("strategy_action_type")=="improve_existing"]
    # Only surface case-level new thesis candidates (NEW_STRATEGY_THESIS_CANDIDATE), not exploratory families
    new_t = list(set(c.get("case_id","") for c in eligible
                     if c.get("decision_grade")=="NEW_STRATEGY_THESIS_CANDIDATE"))
    not_yet=[c.get("case_id","") for c in eligible if c.get("decision_grade")=="NOT_RELIABLE_YET"]
    t=_make_table(doc, ["Action","Items"], [2.0,6.6])
    _add_row(t, ["Keep tracking (repeated sigs):", ", ".join(keep) or "none today"])
    _add_row(t, ["Improve existing (case-level):", ", ".join(imp) or "none today"])
    _add_row(t, ["New thesis case-level (decision_grade=NEW_STRATEGY_THESIS_CANDIDATE):",
                 ", ".join(new_t) or "none today"])
    _add_row(t, ["Not conclude yet:",              ", ".join(not_yet) or "none"])


# ---------------------------------------------------------------------------
# Trap detection
# ---------------------------------------------------------------------------

def _detect_traps(cases):
    traps=[]
    runaway=[c.get("symbol","") for c in cases if c.get("structural_quality")=="runaway_no_base"]
    if runaway:
        traps.append({"trap_code":"RUNAWAY_NO_BASE","description":"Break without compression base",
                      "affected":", ".join(runaway[:5]),"why_risky":"Entry chases, not positions",
                      "caution":"Do not treat as signature candidate"})
    crowd=[c.get("symbol","") for c in cases if c.get("participation_pattern")=="crowd_chase_dominant"
           and c.get("structural_quality") not in ("clean_base_break","repeated_test_then_break")]
    if crowd:
        traps.append({"trap_code":"CROWD_CHASE_DOMINANT","description":"Crowd-driven without clean structure",
                      "affected":", ".join(crowd[:5]),"why_risky":"Reverses quickly without large participant support",
                      "caution":"Require large proxy confirmation"})
    lowpart=[c.get("symbol","") for c in cases if c.get("participation_pattern")=="low_participation_move"]
    if lowpart:
        traps.append({"trap_code":"LOW_PARTICIPATION_MOVE","description":"Minimal flow participation at break",
                      "affected":", ".join(lowpart[:5]),"why_risky":"Low participation breaks often fail",
                      "caution":"Mark descriptive_only until confirmed"})
    weak=[c.get("symbol","") for c in cases if (c.get("proxy_completeness_score") or 0)<40]
    if len(weak)>=3:
        traps.append({"trap_code":"WEAK_PROXY","description":f"{len(weak)} tokens proxy < 40%",
                      "affected":", ".join(weak[:5]),"why_risky":"Unreliable flow classification",
                      "caution":"Treat conclusions as provisional"})
    return traps


def _avg_field(cases, f):
    vs=[]
    for c in cases:
        v=c.get(f)
        if v not in (None,"","None"):
            try: vs.append(float(v))
            except: pass
    return round(sum(vs)/len(vs),4) if vs else None



# ---------------------------------------------------------------------------
# New sections (Decision Bridge — inserted after Section 14)
# ---------------------------------------------------------------------------

_ROLE_COLORS_DB = {
    "repeated_candidate": "70AD47",
    "tracking":           "FFC000",
    "first_observation":  "2E75B6",
    "stale":              "808080",
}

_ANTIPATTERN_WHY = {
    "RUNAWAY_NO_BASE":         "Entry chases price with no base; high reversal risk",
    "DIRTY_BREAK":             "Break quality insufficient; unclear participation",
    "LOW_PARTICIPATION_BREAK": "Minimal flow at break; often fails without follow-through",
    "CROWD_CHASE_DOMINANT":    "Crowd-driven without large participant confirmation; reverses quickly",
}


def _snew_semantic_warning(doc, normalized_ledger_rows):
    warnings = validate_ledger_semantics(normalized_ledger_rows)
    if not warnings:
        return
    _h2(doc, "14a. Ledger Semantic Warnings")
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]; _shade(cell, "FFF2CC")
    cell.paragraphs[0].clear()
    for w in warnings:
        p = cell.add_paragraph(f"\u26a0  {w}")
        for r in p.runs:
            r.font.size = Pt(9); r.bold = True
    doc.add_paragraph()


def _snew_intervention_shortlist(doc, cases, sig_candidates, normalized_ledger_rows):
    _h2(doc, "14b. Strategy Intervention Shortlist")
    eligible = [c for c in cases if c.get("research_eligible_YN") == "Y"]
    shortlist = build_intervention_shortlist(eligible, sig_candidates, normalized_ledger_rows)
    if not shortlist:
        _p(doc, "No intervention candidates today."); doc.add_paragraph(); return
    t = _make_table(doc,
        ["Strategy Family", "Issue Layer", "Cases", "Repeated?", "Cross-Day", "Evidence Source", "Readiness", "Next Action"],
        [1.5, 1.3, 0.5, 0.6, 0.7, 1.3, 1.5, 1.2])
    for row in shortlist:
        _add_row(t, [row.get("strategy_family",""), row.get("issue_layer",""),
            row.get("case_count_today",0), str(row.get("repeated_support_today",0)),
            row.get("cross_day_support",0), row.get("evidence_source",""), row.get("readiness",""),
            row.get("next_action","")])
    doc.add_paragraph()


def _snew_antipattern_board(doc, cases):
    _h2(doc, "14c. Anti-Pattern / Downgrade Board")
    patterns = []
    runaway = [c for c in cases if c.get("structural_quality") == "runaway_no_base"]
    if runaway:
        patterns.append({"code":"RUNAWAY_NO_BASE","affected":", ".join(c.get("symbol","") for c in runaway[:5]),
            "why":_ANTIPATTERN_WHY["RUNAWAY_NO_BASE"],"repeated":"Y" if len(runaway)>=2 else "N"})
    dirty = [c for c in cases if c.get("structural_quality") == "dirty_break"]
    if dirty:
        patterns.append({"code":"DIRTY_BREAK","affected":", ".join(c.get("symbol","") for c in dirty[:5]),
            "why":_ANTIPATTERN_WHY["DIRTY_BREAK"],"repeated":"Y" if len(dirty)>=2 else "N"})
    lowpart = [c for c in cases if c.get("participation_pattern") == "low_participation_move"]
    if lowpart:
        patterns.append({"code":"LOW_PARTICIPATION_BREAK","affected":", ".join(c.get("symbol","") for c in lowpart[:5]),
            "why":_ANTIPATTERN_WHY["LOW_PARTICIPATION_BREAK"],"repeated":"Y" if len(lowpart)>=2 else "N"})
    crowd = [c for c in cases if c.get("participation_pattern") == "crowd_chase_dominant"
             and c.get("structural_quality") not in ("clean_base_break","repeated_test_then_break","exhaustion_spike")]
    if crowd:
        patterns.append({"code":"CROWD_CHASE_DOMINANT","affected":", ".join(c.get("symbol","") for c in crowd[:5]),
            "why":_ANTIPATTERN_WHY["CROWD_CHASE_DOMINANT"],"repeated":"Y" if len(crowd)>=2 else "N"})
    if not patterns:
        _p(doc, "No anti-pattern flags today."); doc.add_paragraph(); return
    t = _make_table(doc, ["Pattern Code","Affected Symbols","Why Downgrade","Repeated Today","Note"],
                    [1.6,1.5,2.4,0.9,1.8])
    for pt in patterns:
        _add_row(t, [pt["code"],pt["affected"],pt["why"],pt["repeated"],"do not use as strategy thesis"])
    doc.add_paragraph()


def _snew_ledger_snapshot(doc, normalized_ledger_rows, research_day):
    _h2(doc, f"14d. Cross-Day Ledger Snapshot (as of {research_day})")
    snapshot = build_ledger_snapshot_for_report(normalized_ledger_rows, research_day)
    if not snapshot:
        _p(doc, "No ledger data available."); doc.add_paragraph(); return
    t = _make_table(doc,
        ["Code","First Seen","Last Seen","Support Days","Recent 7d","Status","Role"],
        [2.0,0.9,0.9,0.9,0.8,1.3,1.4])
    for row in snapshot:
        role = row.get("current_role","")
        r = t.add_row()
        vals = [row.get("signature_candidate_code", row.get("signature_key","")[:25]),
                row.get("first_seen_date",""), row.get("last_seen_date",""),
                row.get("support_days_count",0), row.get("recent_support_days_count",0),
                row.get("latest_validation_status",""), role]
        for i, v in enumerate(vals):
            cell = r.cells[i]; cell.text = str(v) if v is not None else "\u2014"; _font(cell)
            if i == 6:
                color = _ROLE_COLORS_DB.get(role)
                if color:
                    _shade(cell, color)
                    for p in cell.paragraphs:
                        for run in p.runs: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            elif role == "stale":
                _shade(cell, "D9D9D9")
    doc.add_paragraph()


def _snew_promotion_rules(doc):
    _h2(doc, "14e. Readiness Promotion Rules")
    t = _make_table(doc, ["Readiness Level","Meaning"], [2.5,5.8], fill="404040")
    for level, meaning in [
        ("descriptive_only",                  "Not enough repetition or strategy relevance."),
        ("keep_tracking",                      "Interesting but insufficient for intervention."),
        ("old_strategy_improvement_candidate", "Points to existing strategy + identifiable improvement layer."),
        ("new_strategy_thesis_candidate",      "New family candidate — not yet validated, needs multi-day support."),
    ]:
        _add_row(t, [level, meaning])
    _p(doc, "Note: Repeated signatures (Section 10) and case-level theses (Section 13) are independent. Do not conflate.",
       size=9, italic=True)
    doc.add_paragraph()




# ---------------------------------------------------------------------------
# Section 14f — Outcome Horizon Contract Note (static)
# ---------------------------------------------------------------------------

def _snew_outcome_horizon_note(doc):
    _h2(doc, "14f. Outcome Horizon Contract")
    t = _make_table(doc, ["Horizon", "Role", "Usage Rule"], [1.2, 1.8, 5.3], fill="404040")
    _add_row(t, ["1h", "Reaction horizon", "Primary: first directional signal after P2. Used for all family answer contracts."])
    _add_row(t, ["4h", "Tail horizon", "Primary: sustained move / reversal confirmation. Used for resolution label and missed-opportunity assessment."])
    _add_row(t, ["2h", "Diagnostic only", "Optional: interim diagnostic when available. Do NOT mix with 1h or 4h in comparisons without explicit labeling."])
    _p(doc,
       "Do not mix horizons silently. If a board uses 2h data, it must be labeled explicitly. "
       "All family-level measurement boards in this pack use 1h and 4h only unless stated otherwise.",
       size=9, italic=True)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 14g — Measurement Decision Card
# ---------------------------------------------------------------------------

_DECISION_STATE_COLORS = {
    "NO_CHANGE":                      "808080",
    "KEEP_TRACKING":                  "FFC000",
    "PREPARE_HYPOTHESIS":             "2E75B6",
    "READY_FOR_CONTROLLED_VALIDATION": "70AD47",
}


def _snew_measurement_decision_card(doc, cases, sig_candidates, normalized_ledger_rows):
    _h2(doc, "14g. Measurement Decision Card")
    _p(doc,
       "One-family-at-a-time measurement decision. Conservative by design. "
       "Repeated signatures are NOT strategy proof. This card reflects current evidence accumulation only.",
       size=9, italic=True)

    card = build_measurement_decision_card(cases, sig_candidates, normalized_ledger_rows)

    # State row — highlighted
    state = card.get("decision_state", "NO_CHANGE")
    t = _make_table(doc, ["Field", "Value"], [2.2, 6.2], fill="404040")
    state_row = t.add_row()
    state_row.cells[0].text = "Decision State"
    state_row.cells[1].text = state
    _font(state_row.cells[0], bold=True)
    color = _DECISION_STATE_COLORS.get(state, "808080")
    _shade(state_row.cells[1], color)
    for p in state_row.cells[1].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True

    rows = [
        ("Chosen Family",            card.get("chosen_family", "—")),
        ("Chosen Issue Layer",        card.get("chosen_issue_layer", "—")),
        ("Why This Family Now",       card.get("why_this_family_now", "—")),
        ("Expected Upside",           card.get("expected_upside", "—")),
        ("Main Risk / Side Effect",   card.get("main_risk_or_side_effect", "—")),
        ("Evidence Strength",         card.get("evidence_strength_note", "—")),
        ("Why Not Others",            card.get("why_not_others", "—")),
        ("Validation Next Step",      card.get("validation_next_step", "—")),
    ]
    for k, v in rows:
        _add_row(t, [k, v])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 14h — Trusted / Weak / Deferred Summary
# ---------------------------------------------------------------------------

def _snew_trusted_weak_deferred(doc, cases, sig_candidates, selection_context):
    _h2(doc, "14h. Trusted / Weak / Deferred Summary")
    _p(doc,
       "Explicit separation of what is trusted today, what is weak, "
       "and what must be resolved before any strategy action.",
       size=9, italic=True)

    twd = build_trusted_weak_deferred(cases, sig_candidates, selection_context)

    labels = [
        ("Trusted Today",        twd.get("trusted", []),  "70AD47"),
        ("Weak Today",           twd.get("weak", []),     "FFC000"),
        ("Deferred Before Action", twd.get("deferred", []), "C00000"),
    ]
    for label, items, color in labels:
        # Sub-header row as shaded single-cell
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        hdr_cell = tbl.rows[0].cells[0]
        hdr_cell.text = label
        _shade(hdr_cell, color)
        _white_bold(hdr_cell)
        # Items
        for item in items:
            row = tbl.add_row()
            row.cells[0].text = f"• {item}"
            _font(row.cells[0], size=9)
        doc.add_paragraph()

# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------



# ---------------------------------------------------------------------------
# Phase 2A — shared board rendering helper
# ---------------------------------------------------------------------------

_BOARD_H = ["Group", "N", "F1h", "F4h", "A4h", "Dom.Res", "Ready"]
_BOARD_W = [2.3, 0.3, 0.6, 0.6, 0.6, 1.5, 1.3]


def _render_compact_board(doc, title, question, rows, proxy_note=""):
    _p(doc, title, bold=True, size=10)
    _p(doc, f"Q: {question}", size=9, italic=True)
    if proxy_note:
        _p(doc, f"Group: {proxy_note}", size=9, italic=True)
    if not rows:
        _p(doc, "No combinations in current eligible SHORT cases.", size=9)
        doc.add_paragraph()
        return
    t = _make_table(doc, _BOARD_H, _BOARD_W, fill="2E75B6")
    for row in rows:
        _add_row(t, [
            row.get("group_label", "—"),
            row.get("case_count", "—"),
            _fmt(row.get("median_f1h")),
            _fmt(row.get("median_f4h")),
            _fmt(row.get("median_a4h")),
            row.get("dominant_resolution", "—"),
            row.get("readiness_note", "—"),
        ])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2A — 14c2  Short Family Measurement Boards
# ---------------------------------------------------------------------------

def _snew_short_family_boards(doc, short_boards):
    _h2(doc, "14c2. Short Family Measurement Boards")
    _p(doc,
       "Five compact boards for short-side family measurement. "
       "Eligible SHORT cases only. Directional reading — not action-ready at current sample size. "
       "F1h/F4h = median favor pct. A4h = median adverse pct.",
       size=9, italic=True)
    count = short_boards.get("short_eligible_count", 0)
    _p(doc, f"Eligible SHORT cases today: {count}  |  {short_boards.get('note', '')}")
    if count == 0:
        doc.add_paragraph()
        return
    _render_compact_board(
        doc, "B1. Breakdown Quality Board",
        "Which short breakdown structures look promising vs dirty/fake?",
        short_boards.get("b1_breakdown_quality", []),
        "move_class | structural_quality | break_quality_band | resolution_label",
    )
    _render_compact_board(
        doc, "B2. Retest / Reclaim Behavior Board",
        "Are short candidates failing because reclaim is too deep or fail confirmation is weak/late?",
        short_boards.get("b2_retest_reclaim", []),
        "move_class | structural_quality | reclaim_break_4h_YN (proxy) | resolution_label",
    )
    _render_compact_board(
        doc, "B3. Exhaustion / Top Behavior Board",
        "Is exhaustion/top behavior usable or still descriptive-only?",
        short_boards.get("b3_exhaustion_top", []),
        "pre_move_signature | participation_pattern | move_class | resolution_label",
    )
    _render_compact_board(
        doc, "B4. Timing / Staleness Board",
        "Are some shorts weak because they are late/stale rather than structurally invalid?",
        short_boards.get("b4_timing_staleness", []),
        "display_family | timing_band (from time_to_2pct_favor_min) | resolution_label",
    )
    _render_compact_board(
        doc, "B5. Context / Regime Board",
        "Does the same short structure behave differently by regime / breadth / BTC context?",
        short_boards.get("b5_context_regime", []),
        "research_regime | btc_bucket | breadth_bucket | resolution_label",
    )


# ---------------------------------------------------------------------------
# Phase 2A — 14c3  Interaction Board
# ---------------------------------------------------------------------------

def _snew_interaction_board(doc, interaction_data):
    _h2(doc, "14c3. Interaction Board")
    _p(doc,
       "Combined-pattern behavior. Four interaction pairs. Top 5 combinations per pair. "
       "Directional only — not action-ready at current sample size.",
       size=9, italic=True)
    if not interaction_data:
        _p(doc, "No eligible cases for interaction analysis today.")
        doc.add_paragraph()
        return
    _ch = ["Combination", "N", "F1h", "F4h", "A4h", "Dom.Res", "Sample"]
    _cw = [2.3, 0.3, 0.6, 0.6, 0.6, 1.5, 0.9]
    for pair in interaction_data:
        _p(doc, pair.get("interaction_pair", "—"), bold=True, size=9)
        combos = pair.get("top_combinations", [])
        if not combos:
            _p(doc, "No combinations found.", size=9)
            doc.add_paragraph()
            continue
        t = _make_table(doc, _ch, _cw, fill="404040")
        for combo in combos:
            _add_row(t, [
                combo.get("combination_key", "—"),
                combo.get("case_count", "—"),
                _fmt(combo.get("median_f1h")),
                _fmt(combo.get("median_f4h")),
                _fmt(combo.get("median_a4h")),
                combo.get("dominant_resolution", "—"),
                combo.get("sample_note", "—"),
            ])
        _p(doc, pair.get("analyst_note", ""), size=9, italic=True)
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2A — 14f2  Anchor QA / Anchor Audit Summary
# ---------------------------------------------------------------------------

def _snew_anchor_qa(doc, anchor_qa):
    _h2(doc, "14f2. Anchor QA / Anchor Audit Summary")
    _p(doc,
       "Anchor render completeness vs manual audit readiness. "
       "anchor_measurement_readiness cannot exceed DIRECTIONAL_ONLY "
       "without a manual spot-check sample.",
       size=9, italic=True)
    readiness = anchor_qa.get("anchor_measurement_readiness", "NOT_RUN")
    t = _make_table(doc, ["Field", "Value"], [3.2, 4.5], fill="404040")
    r0 = t.add_row()
    r0.cells[0].text = "anchor_measurement_readiness"
    r0.cells[1].text = readiness
    _font(r0.cells[0], bold=True)
    _shade(r0.cells[1], _READINESS_COLORS.get(readiness, "808080"))
    for p in r0.cells[1].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    _add_row(t, ["anchor_rows_expected",           str(anchor_qa.get("anchor_rows_expected", "—"))])
    _add_row(t, ["anchor_rows_rendered",            str(anchor_qa.get("anchor_rows_rendered", "—"))])
    _add_row(t, ["anchor_images_rendered",          str(anchor_qa.get("anchor_images_rendered", "—"))])
    _add_row(t, ["anchor_detect_method",            str(anchor_qa.get("anchor_detect_method", "—"))])
    _add_row(t, ["anchor_conflict_cases_count",     str(anchor_qa.get("anchor_conflict_cases_count", "—"))])
    _add_row(t, ["anchor_fallback_cases_count",     str(anchor_qa.get("anchor_fallback_cases_count", "—"))])
    _add_row(t, ["anchor_conflict_rows_count",      str(anchor_qa.get("anchor_conflict_rows_count", "—"))])
    _add_row(t, ["manual_anchor_audit_sample_size", str(anchor_qa.get("manual_anchor_audit_sample_size", "—"))])
    _add_row(t, ["manual_anchor_audit_pass_rate",   str(anchor_qa.get("manual_anchor_audit_pass_rate", "—"))])
    doc.add_paragraph()
    _p(doc, anchor_qa.get("anchor_audit_note", ""), size=9, italic=True)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2A — 14g2  Measurement Confidence Summary
# ---------------------------------------------------------------------------

def _snew_measurement_confidence(doc, conf):
    _h2(doc, "14g2. Measurement Confidence Summary")
    _p(doc,
       "Explicit confidence frame. Conservative by design. "
       "confidence_band capped at MEDIUM — HIGH not achievable at Phase 2A stage. "
       "Do not use this section to justify action-ready recommendations.",
       size=9, italic=True)
    cb = conf.get("confidence_band", "LOW")
    t = _make_table(doc, ["Field", "Value"], [3.2, 4.5], fill="404040")
    cb_row = t.add_row()
    cb_row.cells[0].text = "confidence_band"
    cb_row.cells[1].text = cb
    _font(cb_row.cells[0], bold=True)
    _shade(cb_row.cells[1], _CONFIDENCE_BAND_COLORS.get(cb, "808080"))
    for p in cb_row.cells[1].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
    _add_row(t, ["decision_sample_size",            str(conf.get("decision_sample_size", "—"))])
    _add_row(t, ["chosen_family_sample_size",        str(conf.get("chosen_family_sample_size", "—"))])
    _add_row(t, ["largest_short_family_sample_size", str(conf.get("largest_short_family_sample_size", "—"))])
    _add_row(t, ["sample_gate_status",
                 conf.get("sample_gate_status", "—")],    hi_col=1, hi_map=_GATE_COLORS)
    _add_row(t, ["bucket_gate_status",
                 conf.get("bucket_gate_status", "—")],    hi_col=1, hi_map=_GATE_COLORS)
    _add_row(t, ["stability_flag",
                 conf.get("stability_flag", "—")],        hi_col=1, hi_map=_STABILITY_COLORS_P2)
    _add_row(t, ["regime_consistency_flag",
                 conf.get("regime_consistency_flag", "—")], hi_col=1, hi_map=_REGIME_COLORS_P2)
    doc.add_paragraph()
    _p(doc, conf.get("confidence_note", ""), size=9, italic=True)
    doc.add_paragraph()
    t2 = _make_table(doc, ["Gate", "Requirement"], [2.0, 5.7], fill="404040")
    _add_row(t2, ["Why not promoted",           conf.get("why_not_promoted", "—")])
    _add_row(t2, ["Next validation requirement", conf.get("next_validation_requirement", "—")])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2C — 14f3  Layer Field Coverage Audit
# ---------------------------------------------------------------------------

_COVERAGE_STATUS_COLORS = {
    "OK":       "70AD47",
    "PARTIAL":  "FFC000",
    "BLOCKING": "C00000",
}

def _snew_layer_field_coverage(doc, layer_audit: dict) -> None:
    _h2(doc, "14f3. Layer Field Coverage Audit")
    _p(doc,
       "Audit of mandatory fields against the Layer 0–8 master spec. "
       "BLOCKING = mandatory field missing from case schema. "
       "PARTIAL = field present but always null. "
       "OK = field present and usable. "
       "Layers 3–6 showing BLOCKING is expected Phase 2C output — not a bug.",
       size=9, italic=True)

    if not layer_audit:
        _p(doc, "Layer audit not available.", size=9, italic=True)
        doc.add_paragraph()
        return

    # Summary card
    t_sum = _make_table(doc, ["Field", "Value"], [3.2, 4.5], fill="404040")
    _add_row(t_sum, ["total_required_fields",    str(layer_audit.get("total_required_fields", "—"))])
    _add_row(t_sum, ["fields_present_count",      str(layer_audit.get("fields_present_count", "—"))])
    _add_row(t_sum, ["fields_missing_count",      str(layer_audit.get("fields_missing_count", "—"))])
    _add_row(t_sum, ["fields_all_null_count",     str(layer_audit.get("fields_all_null_count", "—"))])
    _add_row(t_sum, ["blocking_layers_count",     str(layer_audit.get("blocking_layers_count", "—"))])
    _add_row(t_sum, ["blocking_layers",           str(layer_audit.get("blocking_layers", "—"))])
    doc.add_paragraph()

    # Per-layer compact table
    layer_rows = layer_audit.get("layer_rows", [])
    if layer_rows:
        t = _make_table(doc,
            ["Layer", "Name", "Req", "Present", "Missing", "Null", "Status"],
            [0.4, 2.2, 0.5, 0.6, 0.6, 0.5, 0.9])
        for row in layer_rows:
            status = row.get("coverage_status", "—")
            tr = t.add_row()
            vals = [
                str(row.get("layer_id", "—")),
                row.get("layer_name", "—")[:30],
                str(row.get("required_fields_count", "—")),
                str(row.get("present_fields_count", "—")),
                str(row.get("missing_fields_count", "—")),
                str(row.get("all_null_fields_count", "—")),
                status,
            ]
            for i, v in enumerate(vals):
                cell = tr.cells[i]
                cell.text = v
                _font(cell)
                if i == 6:
                    color = _COVERAGE_STATUS_COLORS.get(status)
                    if color:
                        _shade(cell, color)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2B — 14g3  Multi-Day Validation Snapshot (Phase 2C updated)
# ---------------------------------------------------------------------------

def _snew_multiday_validation_snapshot(
    doc,
    validation_summary: dict,
    family_history_snapshot: Optional[dict] = None,  # accepted but NOT used — Phase 2B is independent
) -> None:
    _h2(doc, "14g3. Multi-Day Validation Snapshot")
    # Phase 2B is today-only conservative promotion frame.
    # Phase 2C family history is in section 14g3b and is completely independent.
    # Do NOT override Phase 2B fields with Phase 2C data.
    _p(doc,
       "Phase 2B — today-only conservative promotion frame. "
       "SHORT-side, research_eligible_YN == Y. "
       "Phase 2C family history (multiday) is shown separately in 14g3b. "
       "PREPARE_HYPOTHESIS blocked — Phase 2E gates required "
       "(anchor QA, unseen-day validation, sample thresholds). "
       "Bucket-ready: >= 10. Tracking-grade: >= 20. Recommendation: >= 50.",
       size=9, italic=True)

    if not validation_summary:
        _p(doc, "No validation summary available.", size=9, italic=True)
        doc.add_paragraph()
        return

    promotion_state = validation_summary.get("promotion_state", "DESCRIPTIVE_ONLY")

    t = _make_table(doc, ["Field", "Value"], [3.2, 4.5], fill="404040")
    ps_row = t.add_row()
    ps_row.cells[0].text = "promotion_state"
    ps_row.cells[1].text = promotion_state
    _font(ps_row.cells[0], bold=True)
    _shade(ps_row.cells[1], _PROMOTION_COLORS.get(promotion_state, "808080"))
    for p in ps_row.cells[1].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    _add_row(t, ["Top Family (today only)",
                 str(validation_summary.get("top_candidate_family", "—"))])
    _add_row(t, ["Case Count (today only)",
                 str(validation_summary.get("_top_case_count", "—"))])
    _add_row(t, ["Ledger Research Days (global context only)",
                 str(validation_summary.get("ledger_research_days_global_context", "—"))])
    _add_row(t, ["families_at_bucket_ready_or_above (>= 10)",
                 str(validation_summary.get("families_at_bucket_ready_or_above", "—"))])
    _add_row(t, ["families_at_tracking_grade_or_above (>= 20)",
                 str(validation_summary.get("families_at_tracking_grade_or_above", "—"))])
    _add_row(t, ["families_at_recommendation_grade_or_above (>= 50)",
                 str(validation_summary.get("families_at_recommendation_grade_or_above", "—"))])
    doc.add_paragraph()

    t2 = _make_table(doc, ["Item", "Detail"], [2.0, 5.7], fill="404040")
    _add_row(t2, ["promotion_reason",     str(validation_summary.get("promotion_reason", "—"))])
    _add_row(t2, ["blocking_reason",      str(validation_summary.get("blocking_reason", "—"))])
    _add_row(t2, ["validation_next_step", str(validation_summary.get("validation_next_step", "—"))])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2C — 14g3b  Family History Compact (inserted after Phase 2B snapshot)
# ---------------------------------------------------------------------------

def _snew_phase2c_family_history_compact(
    doc,
    family_history_snapshot: Optional[dict],
) -> None:
    """Compact Phase 2C family history block — after 14g3, before 14g4."""
    _h2(doc, "14g3b. Phase 2C Family History (compact)")
    _p(doc,
       "Multiday SHORT family history from daily_case_dataset CSVs — NOT from ledger. "
       "PREPARE_HYPOTHESIS remains blocked — Phase 2E gates required "
       "(anchor QA, unseen-day validation, sample thresholds).",
       size=9, italic=True)

    if not family_history_snapshot:
        _p(doc, "Family history snapshot not available.", size=9, italic=True)
        doc.add_paragraph()
        return

    fhs = family_history_snapshot
    t = _make_table(doc, ["Field", "Value"], [3.2, 4.5], fill="404040")
    _add_row(t, ["top_candidate_family",          str(fhs.get("top_candidate_family", "—"))])
    _add_row(t, ["top_family_case_count_multiday", str(fhs.get("top_family_case_count_multiday", "—"))])
    _add_row(t, ["top_family_days_count",          str(fhs.get("top_family_days_count", "—"))])

    hist_status = fhs.get("family_history_status", "not_available_yet")
    hs_row = t.add_row()
    hs_row.cells[0].text = "family_history_status"
    hs_row.cells[1].text = hist_status
    _font(hs_row.cells[0])
    _HIST_STATUS_COLORS = {
        "available":       "70AD47",
        "early_tracking":  "FFC000",
        "not_available_yet": "808080",
    }
    color = _HIST_STATUS_COLORS.get(hist_status)
    if color:
        _shade(hs_row.cells[1], color)
        for p in hs_row.cells[1].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _add_row(t, ["families_with_history_count",   str(fhs.get("families_with_history_count", "—"))])
    _add_row(t, ["historical_case_days_loaded",    str(fhs.get("historical_case_days_loaded", "—"))])
    doc.add_paragraph()
    _p(doc, fhs.get("history_note", ""), size=8, italic=True)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2D — 14g4  Family Validation Snapshot
# ---------------------------------------------------------------------------

# Color maps for Phase 2D
_BUCKET_STATUS_COLORS_2D = {
    "bucket_ready":         "70AD47",
    "tracking_grade":       "2E75B6",
    "recommendation_grade": "375623",
    "not_enough_sample":    "808080",
}
_STABILITY_COLORS_2D = {
    "STABLE":       "70AD47",
    "EARLY_SIGNAL": "FFC000",
    "UNSTABLE":     "C00000",
    "INSUFFICIENT": "808080",
}
_RECOMMENDATION_COLORS_2D = {
    "keep_tracking":     "2E75B6",
    "descriptive_only":  "808080",
    "blocked":           "C00000",
    "no_change":         "808080",
}

_P2E_STATE_COLORS = {
    "READY_FOR_CONTROLLED_VALIDATION":     "375623",
    "PREPARE_HYPOTHESIS":                  "538135",
    "NOT_READY_FOR_CONTROLLED_VALIDATION": "C55A11",
    "KEEP_TRACKING":                       "2E75B6",
    "DESCRIPTIVE_ONLY":                    "808080",
}

_P2E_UNSEEN_COLORS = {
    "NOT_READY_no_holdout_data":             "808080",
    "NOT_READY_insufficient_holdout_sample": "C55A11",
    "PARTIAL_holdout_available":             "2E75B6",
}


def _snew_phase2d_validation_snapshot(
    doc,
    family_validation_stats: Optional[list],
) -> None:
    """Phase 2D — Section 14g4: Family Validation Snapshot."""
    _h2(doc, "14g4. Phase 2D Family Validation Snapshot")
    _p(doc,
       "Phase 2D — statistical validation per SHORT family using multiday case pool. "
       "Bootstrap CI (95%) from stdlib random. KS/t-test via optional scipy. "
       "Conservative: blocked / not_enough_sample outputs are expected at current sample sizes. "
       "PREPARE_HYPOTHESIS is NOT unlocked by Phase 2D stats alone.",
       size=9, italic=True)

    if not family_validation_stats:
        _p(doc,
           "No Phase 2D validation stats available. "
           "Run after historical daily_case_dataset CSVs exist.",
           size=9, italic=True)
        doc.add_paragraph()
        return

    t = _make_table(doc,
        ["Family", "N multi", "Days", "Med F1h", "Med F4h", "Med A4h",
         "CI Low", "CI High", "Bucket", "Stability", "KS", "t-test"],
        [1.8, 0.55, 0.45, 0.6, 0.6, 0.6, 0.55, 0.55, 0.75, 0.75, 0.9, 1.0])

    for row in family_validation_stats:
        bk = row.get("bucket_ready_status", "—")
        st = row.get("stability_flag", "—")
        is_actionable = row.get("is_actionable_family", True)
        tr = t.add_row()
        vals = [
            row.get("display_family", "—"),
            row.get("case_count_multiday", "—"),
            row.get("family_days_count", "—"),
            _fmt(row.get("median_f1h")),
            _fmt(row.get("median_f4h")),
            _fmt(row.get("median_a4h")),
            _fmt(row.get("bootstrap_ci_low")),
            _fmt(row.get("bootstrap_ci_high")),
            bk,
            st,
            str(row.get("ks_test_status_or_na", "—")),
            str(row.get("ttest_status_or_na", "—")),
        ]
        for i, v in enumerate(vals):
            cell = tr.cells[i]
            cell.text = str(v) if v is not None else "—"
            _font(cell)
            if i == 8:
                # Population buckets (unclassified) get grey regardless of grade
                color = "808080" if not is_actionable else _BUCKET_STATUS_COLORS_2D.get(str(v))
                if color:
                    _shade(cell, color)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i == 9:
                color = _STABILITY_COLORS_2D.get(str(v))
                if color:
                    _shade(cell, color)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # Statistical readiness notes (one paragraph per family)
    for row in family_validation_stats:
        note = row.get("statistical_readiness_note", "")
        if note:
            _p(doc, f"[{row.get('display_family','?')}] {note}", size=8, italic=True)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Phase 2D — 14g5  Family Answer Contracts
# ---------------------------------------------------------------------------

def _snew_phase2d_answer_contracts(
    doc,
    family_answer_contracts: Optional[list],
) -> None:
    """Phase 2D — Section 14g5: Family Answer Contracts (6 question families)."""
    _h2(doc, "14g5. Phase 2D Family Answer Contracts")
    _p(doc,
       "One contract per question family: Exhaustion / Breakdown / Retest fail / "
       "Timing / Invalidation / Context. "
       "Blocked when required Layer fields are missing at source. "
       "Does NOT write live-rule instructions. "
       "recommendation_state stays conservative — see validation_next_step to unblock.",
       size=9, italic=True)

    if not family_answer_contracts:
        _p(doc,
           "No Phase 2D answer contracts available. "
           "Run after historical case datasets exist.",
           size=9, italic=True)
        doc.add_paragraph()
        return

    for contract in family_answer_contracts:
        fname = contract.get("family_name", "—")
        rstate = contract.get("recommendation_state", "blocked")

        _p(doc, fname, bold=True, size=10)

        t = _make_table(doc, ["Field", "Detail"], [2.2, 5.5], fill="404040")

        # recommendation_state gets color treatment
        rs_row = t.add_row()
        rs_row.cells[0].text = "recommendation_state"
        rs_row.cells[1].text = rstate
        _font(rs_row.cells[0], bold=True)
        color = _RECOMMENDATION_COLORS_2D.get(rstate)
        if color:
            _shade(rs_row.cells[1], color)
            for p in rs_row.cells[1].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True

        _add_row(t, ["strategy_question",
                     str(contract.get("strategy_question", "—"))[:200]])
        _add_row(t, ["threshold_direction",
                     str(contract.get("threshold_direction", "—"))])
        _add_row(t, ["evidence_summary",
                     str(contract.get("evidence_summary", "—"))[:200]])
        _add_row(t, ["confidence_summary",
                     str(contract.get("confidence_summary", "—"))])
        _add_row(t, ["blocking_reason",
                     str(contract.get("blocking_reason", "—"))[:180]])
        _add_row(t, ["likely_side_effect",
                     str(contract.get("likely_side_effect", "—"))[:180]])
        _add_row(t, ["validation_next_step",
                     str(contract.get("validation_next_step", "—"))[:180]])
        doc.add_paragraph()




def _trunc(text, limit: int = 130) -> str:
    """Clean truncation with ellipsis for Phase 2E display fields."""
    s = str(text) if text is not None else "—"
    if len(s) <= limit:
        return s
    # Break at word boundary when possible
    cut = s[:limit - 1].rsplit(" ", 1)[0]
    return cut + "…"


def _snew_phase2e_card(
    doc,
    controlled_validation_state: Optional[list],
    unseen_day_summary: Optional[list],
) -> None:
    """
    Phase 2E — Controlled validation gate card.
    Renderer only: consumes producer outputs, does not invent gate logic.
    """
    _h2(doc, "14g6. Phase 2E — Controlled Validation Gate")
    _p(doc,
       "Gate-building phase. Expected output: conservative for all families. "
       "PREPARE_HYPOTHESIS / READY_FOR_CONTROLLED_VALIDATION require all hard gates. "
       "Conservative NOT_READY / KEEP_TRACKING / DESCRIPTIVE_ONLY is the correct result "
       "at current evidence level.",
       size=9, italic=True)

    if not controlled_validation_state:
        _p(doc,
           "No Phase 2E controlled validation state available. "
           "Requires Phase 2C family history + Phase 2D validation stats.",
           size=9, italic=True)
        doc.add_paragraph()
    else:
        _p(doc, "Controlled Validation State — one row per display_family:", bold=True, size=9)
        t = _make_table(
            doc,
            ["Family", "State", "N multi", "Days", "Anchor Gate", "Blocker Summary"],
            [1.8, 2.0, 0.6, 0.5, 1.2, 2.3],
        )
        for row in controlled_validation_state:
            state = row.get("controlled_validation_state", "—")
            tr = t.add_row()
            vals = [
                row.get("display_family", "—"),
                state,
                row.get("case_count_multiday", "—"),
                row.get("family_days_count", "—"),
                row.get("anchor_qa_gate", "—"),
                _trunc(row.get("promotion_blocker_summary", "—"), 130),
            ]
            for i, v in enumerate(vals):
                cell = tr.cells[i]
                cell.text = str(v) if v is not None else "—"
                _font(cell, size=8)
                if i == 1:
                    color = _P2E_STATE_COLORS.get(str(v))
                    if color:
                        _shade(cell, color)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph()

    if not unseen_day_summary:
        _p(doc, "No unseen-day validation summary available.", size=9, italic=True)
        doc.add_paragraph()
        return

    _p(doc, "Unseen-Day / Holdout Validation Summary:", bold=True, size=9)
    t2 = _make_table(
        doc,
        ["Family", "Unseen Status", "Holdout N", "Days", "Summary"],
        [1.8, 2.0, 0.65, 0.5, 3.45],
    )
    for row in unseen_day_summary:
        status = row.get("unseen_validation_status", "—")
        tr = t2.add_row()
        vals = [
            row.get("display_family", "—"),
            status,
            row.get("holdout_case_count", 0),
            row.get("holdout_days_used", 0),
            _trunc(row.get("holdout_result_summary", "—"), 130),
        ]
        for i, v in enumerate(vals):
            cell = tr.cells[i]
            cell.text = str(v) if v is not None else "—"
            _font(cell, size=8)
            if i == 1:
                color = _P2E_UNSEEN_COLORS.get(str(v))
                if color:
                    _shade(cell, color)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()
    _p(doc,
       "> Unseen-day validation not yet run. "
       "Holdout data required before any promotion claim is valid.",
       size=8, italic=True)
    doc.add_paragraph()

def build_docx_pack(
    research_day, selection_context, cases, anchor_rows, signature_candidates,
    output_path, image_results_all=None,
    layer_audit: Optional[dict] = None,
    family_history_snapshot: Optional[dict] = None,
    p2d_family_validation_stats: Optional[list] = None,
    p2d_family_answer_contracts: Optional[list] = None,
    p2e_controlled_validation_state: Optional[list] = None,
    p2e_unseen_day_summary: Optional[list] = None,
) -> str:
    if image_results_all is None: image_results_all={}
    normalized_ledger_rows = load_and_normalize_ledger_rows(research_day, window_days=7, as_of_day=research_day)

    # Phase 2A pre-computations
    anchor_qa        = build_anchor_qa_summary(cases, anchor_rows)
    short_boards     = build_short_family_boards(cases)
    interaction_data = build_interaction_board(cases)
    conf_summary     = build_measurement_confidence_summary(
        cases, signature_candidates, normalized_ledger_rows)

    # Phase 2B pre-computations
    _ledger_days_context = len(set(
        r.get("research_day", "") for r in normalized_ledger_rows
        if r.get("research_day")
    ))
    multiday_fam_stats    = build_multiday_family_stats(cases, normalized_ledger_rows)
    multiday_inter_stats  = build_multiday_interaction_stats(cases)
    validation_summary    = build_controlled_validation_summary(
        multiday_fam_stats, multiday_inter_stats, _ledger_days_context)
    # Attach top case_count aligned with top_candidate_family (named family, not unclassified)
    if multiday_fam_stats:
        _top_fam = validation_summary.get("top_candidate_family", "")
        _top_row = next((r for r in multiday_fam_stats if r.get("display_family") == _top_fam), None)
        validation_summary["_top_case_count"] = (
            _top_row.get("case_count", "—") if _top_row else "—"
        )

    doc=Document()

    _h1(doc, "Daily Top Movers Research Pack")
    _p(doc, f"Research Day: {research_day}", bold=True)
    _p(doc, f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    _p(doc, f"Regime: {selection_context.get('research_regime','—')}")
    _p(doc, f"BTC 24h: {_fmt(selection_context.get('btc_24h_change_pct'),2)}%  |  Alt Breadth: {_fmt(selection_context.get('alt_breadth_pct'),1)}%")
    eligible_n=sum(1 for c in cases if c.get("research_eligible_YN")=="Y")
    _p(doc, f"Cases: {len(cases)}  |  Eligible: {eligible_n}  |  Signatures (repeated): {len(signature_candidates)}")
    doc.add_page_break()

    _s1_data_quality(doc, cases)
    _s2_integrity(doc, cases, image_results_all)
    _s3_exec_summary(doc, cases, selection_context, signature_candidates)
    _s4_selection_board(doc, cases)
    _s5_move_archetype(doc, cases)
    _s6_pre_move_sig(doc, cases)
    _s7_participation(doc, cases)
    _s8_structural_quality(doc, cases)
    _s9_outcome_quality(doc, cases)
    _s10_signature_evidence(doc, signature_candidates)
    _s11_strategy_mapping(doc, cases)
    _s12_new_strategy(doc, cases)
    _s13_decision_grade(doc, cases)
    _s14_trap_caution(doc, cases)
    _snew_semantic_warning(doc, normalized_ledger_rows)
    _snew_intervention_shortlist(doc, cases, signature_candidates, normalized_ledger_rows)
    _snew_antipattern_board(doc, cases)
    _snew_short_family_boards(doc, short_boards)
    _snew_interaction_board(doc, interaction_data)
    _snew_ledger_snapshot(doc, normalized_ledger_rows, research_day)
    _snew_promotion_rules(doc)
    _snew_outcome_horizon_note(doc)
    _snew_anchor_qa(doc, anchor_qa)
    _snew_layer_field_coverage(doc, layer_audit or {})
    _snew_measurement_decision_card(doc, cases, signature_candidates, normalized_ledger_rows)
    _snew_measurement_confidence(doc, conf_summary)
    _snew_multiday_validation_snapshot(doc, validation_summary, family_history_snapshot)
    _snew_phase2c_family_history_compact(doc, family_history_snapshot)
    _snew_phase2d_validation_snapshot(doc, p2d_family_validation_stats)
    _snew_phase2d_answer_contracts(doc, p2d_family_answer_contracts)
    _snew_phase2e_card(doc, p2e_controlled_validation_state, p2e_unseen_day_summary)
    _snew_trusted_weak_deferred(doc, cases, signature_candidates, selection_context)
    _s15_review_queue(doc, cases)
    _s16_case_registry(doc, cases)
    _s17_case_appendix(doc, cases, anchor_rows, research_day, image_results_all)
    _s18_footer(doc, cases, signature_candidates)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

#!/usr/bin/env python3
"""
sim_test_suite.py — Regression suite: report semantics + dispatch trace
========================================================================
Runs entirely offline. No Binance API. No market signals needed.

Usage:
  python3 sim_test_suite.py                    # run all tests
  python3 sim_test_suite.py --report-only      # skip unit tests, only build DOCX
  python3 sim_test_suite.py --no-report        # skip DOCX build, only unit tests
  python3 sim_test_suite.py --keep             # keep temp files after run

Archetypes (14 total):
  INVALIDATEDUSDT   INVALIDATED, dispatch=not_routed            (Phase 5A)
  EXPIREDUSDT       EXPIRED_WAIT, dispatch=not_routed
  REJECTEDUSDT      REJECTED_SCORE, dispatch=not_routed
  MAINSIGUSDT       CONFIRMED+SENT, dispatch=MAIN_SIGNAL
  WATCHLISTUSDT     CONFIRMED, dispatch=WATCHLIST
  NOSENDUSDT        CONFIRMED, dispatch=NO_SEND
  PENDINGUSDT       PENDING, dispatch=not_evaluated             (correct)
  SEMBROKENUSDT     CONFIRMED, missing confirmed_ts_ms          (semantic broken)
  SENTBROKENUSDT    CONFIRMED+SENT, missing sent_ts_ms          (NEW — sent_no_ts)
  CLOSEBRKNUSDT     INVALIDATED, closed_ts_ms ok, close_reason=""  (NEW — broken close)
  FALLBACK4HUSDT    INVALIDATED, fallback_4h_snapshot            (NEW — valid close)
  NOTDUEUSDT        CONFIRMED, not_due_yet                      (no fake close ts)
  TRUECLOSEUSDT     INVALIDATED, true_close, KILL_CORRECT
  WAITSHORTUSD      EXPIRED_WAIT, true_close, WAIT_TOO_SHORT_CANDIDATE

Report value assertions (stronger than section-existence checks):
  Section B  closed_no_reason == 1  (only CLOSEBRKNUSDT, not NOTDUEUSDT/FALLBACK4H)
  Section B  sent_no_ts       == 1  (only SENTBROKENUSDT)
  Section B  health banner    == ISSUES FOUND
  Section H  dispatch_action populated > 0 (not_routed counts, Phase 5A)
  Section H  dispatch status  is explicit (N/A, OK, or GAPS — never blank)
  Footer     next_action      == Fix semantic health first
  Footer     NOT "Data clean" when broken rows present
"""

from __future__ import annotations

import argparse
import csv
import shutil
import subprocess
import sys
import tempfile
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional

# ---------------------------------------------------------------------------
# Terminal colours
# ---------------------------------------------------------------------------
GREEN  = "\033[92m"
RED    = "\033[91m"
CYAN   = "\033[96m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

def ok(msg):   print(f"  {GREEN}✓ PASS{RESET}  {msg}")
def fail(msg): print(f"  {RED}✗ FAIL{RESET}  {msg}")
def info(msg): print(f"  {CYAN}ℹ{RESET}      {msg}")

# ---------------------------------------------------------------------------
# Synthetic case factory
# ---------------------------------------------------------------------------
NOW_MS   = int(time.time() * 1000)
DAY      = datetime.now(timezone.utc).strftime("%Y-%m-%d")
CLOSE_MS = NOW_MS - 3 * 3600 * 1000   # 3h ago → 2h outcome available

BASE: Dict = dict(
    setup_id="", created_ts_ms=NOW_MS, signal_open_time=NOW_MS,
    symbol="TESTUSDT", side="LONG", score=80.0, confidence=0.85,
    reason="sim_reason", breakout_level=1.0, signal_price=1.0,
    signal_high=1.01, signal_low=0.99, oi_jump_pct=2.0,
    funding_pct=0.01, vol_ratio=1.8,
    strategy="long_breakout_retest", market_regime="UPTREND",
    btc_price=83000.0, btc_24h_change_pct=1.0, btc_4h_change_pct=0.3,
    btc_1h_change_pct=0.1, btc_24h_range_pct=3.0, btc_4h_range_pct=1.0,
    alt_market_breadth_pct=55.0, btc_regime="bullish",
    score_oi=10.0, score_exhaustion=0.0, score_breakout=30.0,
    score_retest=40.0, reason_tags="sim;test",
    status="PENDING", close_reason="", bars_waited=0,
    closed_ts_ms="", send_decision="", skip_reason="",
    is_confirmed="N", confirmed_ts_ms="",
    is_sent_signal="N", sent_ts_ms="",
    review_eligible="Y", review_exclusion_reason="",
    semantic_consistency="Y", semantic_issue="",
    close_anchor_time_ms="", close_capture_basis="",
    future_1h_max_favor_pct="", future_1h_max_adverse_pct="",
    future_2h_max_favor_pct="", future_2h_max_adverse_pct="",
    future_4h_max_favor_pct="", future_4h_max_adverse_pct="",
    reclaim_breakout_2h_YN="", reclaim_breakout_4h_YN="",
    outcome_1h_available="N", outcome_2h_available="N",
    outcome_1h_summary="", outcome_2h_summary="",
    post_close_outcome_notes="", outcome_conclusion_code="",
    entry_feasible_YN="", entry_feasible_window_minutes="",
    entry_slippage_pct="", entry_execution_note="",
    time_to_max_favor_minutes="", time_to_max_adverse_minutes="",
    close_trigger_detail="", confirm_fail_detail="", invalidation_detail="",
    regret_valid_YN="", regret_filter_reason="",
    regime_label="trend_continuation_friendly",
    regime_fit_for_strategy="HIGH",
    setup_quality_band="MEDIUM", delivery_band="TRADABLE",
    veto_reason_code="not_evaluated",
    dispatch_action="not_evaluated",
    dispatch_confidence_band="not_evaluated",
    dispatch_reason="not_evaluated",
)

_idx = [0]
def make(**overrides) -> Dict:
    _idx[0] += 1
    uid = f"SIM-{_idx[0]:03d}-{NOW_MS}"
    row = dict(BASE)
    row["pending_id"] = uid
    row["setup_id"]   = uid
    row["symbol"]     = overrides.pop("symbol", f"SIM{_idx[0]:03d}USDT")
    row.update(overrides)
    return row


def build_cases() -> List[Dict]:
    return [
        # ── dispatch: non-CONFIRMED terminal → not_routed (Phase 5A fix) ──
        make(symbol="INVALIDATEDUSDT", side="SHORT",
             strategy="short_exhaustion_retest",
             status="INVALIDATED", close_reason="retest invalidated",
             closed_ts_ms=str(NOW_MS - 60_000),
             close_anchor_time_ms=str(NOW_MS - 60_000),
             close_capture_basis="true_close",
             invalidation_detail="retest invalidated",
             dispatch_action="not_routed",
             dispatch_confidence_band="none",
             dispatch_reason="closed_before_dispatch",
             outcome_1h_available="N", outcome_2h_available="N",
             review_eligible="N", review_exclusion_reason="outcome_2h_unavailable"),

        make(symbol="EXPIREDUSDT",
             status="EXPIRED_WAIT", close_reason="retest wait expired",
             closed_ts_ms=str(NOW_MS - 120_000),
             close_anchor_time_ms=str(NOW_MS - 120_000),
             close_capture_basis="true_close",
             confirm_fail_detail="timeout_no_followthrough",
             dispatch_action="not_routed",
             dispatch_confidence_band="none",
             dispatch_reason="closed_before_dispatch",
             outcome_1h_available="N", outcome_2h_available="N",
             review_eligible="N", review_exclusion_reason="outcome_2h_unavailable"),

        make(symbol="REJECTEDUSDT",
             status="REJECTED_SCORE", close_reason="score below min_send 70",
             closed_ts_ms=str(NOW_MS - 90_000),
             close_anchor_time_ms=str(NOW_MS - 90_000),
             close_capture_basis="true_close",
             dispatch_action="not_routed",
             dispatch_confidence_band="none",
             dispatch_reason="closed_before_dispatch",
             outcome_1h_available="N", outcome_2h_available="N",
             review_eligible="N", review_exclusion_reason="outcome_2h_unavailable"),

        # ── dispatch: CONFIRMED paths ──
        make(symbol="MAINSIGUSDT",
             status="CONFIRMED", close_reason="signal confirmed",
             is_confirmed="Y", confirmed_ts_ms=str(NOW_MS - 1_800_000),
             is_sent_signal="Y", sent_ts_ms=str(NOW_MS - 1_800_000),
             send_decision="SENT", close_capture_basis="not_due_yet",
             dispatch_action="MAIN_SIGNAL",
             dispatch_confidence_band="HIGH",
             dispatch_reason="rank=0_top_n=2",
             semantic_consistency="Y",
             outcome_1h_available="N", outcome_2h_available="N",
             post_close_outcome_notes="not_enough_post_close_time"),

        make(symbol="WATCHLISTUSDT", side="SHORT",
             strategy="short_exhaustion_retest",
             status="CONFIRMED", close_reason="signal confirmed",
             is_confirmed="Y", confirmed_ts_ms=str(NOW_MS - 1_800_000),
             send_decision="WATCHLIST", skip_reason="dispatch_watchlist",
             close_capture_basis="not_due_yet",
             dispatch_action="WATCHLIST",
             dispatch_confidence_band="MEDIUM",
             dispatch_reason="rank=2_top_n=2",
             semantic_consistency="Y",
             outcome_1h_available="N", outcome_2h_available="N"),

        make(symbol="NOSENDUSDT",
             status="CONFIRMED", close_reason="signal confirmed",
             is_confirmed="Y", confirmed_ts_ms=str(NOW_MS - 1_800_000),
             send_decision="NO_SEND", skip_reason="dispatch_no_send",
             close_capture_basis="not_due_yet",
             dispatch_action="NO_SEND",
             dispatch_confidence_band="LOW",
             dispatch_reason="score_below_floor",
             semantic_consistency="Y",
             outcome_1h_available="N", outcome_2h_available="N"),

        make(symbol="PENDINGUSDT",
             status="PENDING",
             dispatch_action="not_evaluated",  # correct — not yet dispatched
             dispatch_confidence_band="not_evaluated",
             dispatch_reason="not_evaluated"),

        # ── semantic broken: missing confirmed_ts_ms ──
        make(symbol="SEMBROKENUSDT",
             status="CONFIRMED",
             is_confirmed="Y", confirmed_ts_ms="",   # intentionally missing
             dispatch_action="not_evaluated",
             semantic_consistency="N",
             semantic_issue="confirmed_without_confirmed_ts",
             review_eligible="N",
             review_exclusion_reason="confirmed_without_confirmed_ts"),

        # ── NEW: broken sent path — sent but missing sent_ts_ms ──
        make(symbol="SENTBROKENUSDT",
             status="CONFIRMED", close_reason="signal confirmed",
             is_confirmed="Y", confirmed_ts_ms=str(NOW_MS - 1_800_000),
             is_sent_signal="Y", sent_ts_ms="",      # intentionally missing
             send_decision="SENT",
             close_capture_basis="not_due_yet",
             dispatch_action="MAIN_SIGNAL",
             dispatch_confidence_band="HIGH",
             dispatch_reason="rank=0_top_n=2",
             semantic_consistency="N",
             semantic_issue="sent_without_sent_ts",
             review_eligible="N",
             review_exclusion_reason="sent_without_sent_ts"),

        # ── NEW: broken close — terminal + closed_ts_ms present + close_reason="" ──
        # Only this should increment closed_no_reason counter
        make(symbol="CLOSEBRKNUSDT",
             status="INVALIDATED", close_reason="",  # intentionally missing
             closed_ts_ms=str(NOW_MS - 180_000),
             close_anchor_time_ms=str(NOW_MS - 180_000),
             close_capture_basis="true_close",
             dispatch_action="not_routed",
             dispatch_confidence_band="none",
             dispatch_reason="closed_before_dispatch",
             outcome_1h_available="N", outcome_2h_available="N",
             review_eligible="N", review_exclusion_reason="outcome_2h_unavailable"),

        # ── NEW: fallback_4h_snapshot — valid close, must NOT count as broken ──
        make(symbol="FALLBACK4HUSDT",
             status="INVALIDATED",
             close_reason="fallback_case_close_after_4h",  # present — not broken
             closed_ts_ms=str(NOW_MS - 4 * 3600_000),
             close_anchor_time_ms=str(NOW_MS - 4 * 3600_000),
             close_capture_basis="fallback_4h_snapshot",
             dispatch_action="not_routed",
             dispatch_confidence_band="none",
             dispatch_reason="closed_before_dispatch",
             outcome_1h_available="N", outcome_2h_available="N",
             review_eligible="N", review_exclusion_reason="outcome_2h_unavailable"),

        # ── close semantics: not_due_yet → no fake close timestamp ──
        make(symbol="NOTDUEUSDT",
             status="CONFIRMED", close_reason="signal confirmed",
             is_confirmed="Y", confirmed_ts_ms=str(NOW_MS - 900_000),
             close_capture_basis="not_due_yet",
             closed_ts_ms="", close_anchor_time_ms="",  # no fake ts
             dispatch_action="MAIN_SIGNAL",
             dispatch_confidence_band="HIGH",
             dispatch_reason="rank=0_top_n=2",
             semantic_consistency="Y",
             outcome_1h_available="N", outcome_2h_available="N"),

        # ── outcome: true_close + KILL_CORRECT ──
        make(symbol="TRUECLOSEUSDT",
             status="INVALIDATED", close_reason="stop_loss_hit",
             closed_ts_ms=str(CLOSE_MS),
             close_anchor_time_ms=str(CLOSE_MS),
             close_capture_basis="true_close",
             invalidation_detail="stop_loss_hit",
             dispatch_action="not_routed",
             dispatch_confidence_band="none",
             dispatch_reason="closed_before_dispatch",
             outcome_1h_available="Y", outcome_2h_available="Y",
             outcome_conclusion_code="KILL_CORRECT",
             future_1h_max_favor_pct="0.8000",  future_1h_max_adverse_pct="2.1000",
             future_2h_max_favor_pct="1.1000",  future_2h_max_adverse_pct="2.8000",
             future_4h_max_favor_pct="1.5000",  future_4h_max_adverse_pct="3.2000",
             reclaim_breakout_2h_YN="N", reclaim_breakout_4h_YN="N",
             entry_feasible_YN="N", entry_feasible_window_minutes="10",
             regret_valid_YN="N", regret_filter_reason="entry_not_feasible",
             review_eligible="Y"),

        # ── outcome: true_close + WAIT_TOO_SHORT_CANDIDATE ──
        make(symbol="WAITSHORTUSD",
             status="EXPIRED_WAIT", close_reason="retest wait expired",
             closed_ts_ms=str(CLOSE_MS),
             close_anchor_time_ms=str(CLOSE_MS),
             close_capture_basis="true_close",
             confirm_fail_detail="timeout_no_followthrough",
             dispatch_action="not_routed",
             dispatch_confidence_band="none",
             dispatch_reason="closed_before_dispatch",
             outcome_1h_available="Y", outcome_2h_available="Y",
             outcome_conclusion_code="WAIT_TOO_SHORT_CANDIDATE",
             future_1h_max_favor_pct="2.1000",  future_1h_max_adverse_pct="0.4000",
             future_2h_max_favor_pct="3.8000",  future_2h_max_adverse_pct="0.6000",
             future_4h_max_favor_pct="5.2000",  future_4h_max_adverse_pct="0.8000",
             reclaim_breakout_2h_YN="Y", reclaim_breakout_4h_YN="Y",
             entry_feasible_YN="Y", entry_feasible_window_minutes="10",
             regret_valid_YN="Y", regret_filter_reason="valid_missed_continuation",
             review_eligible="Y"),
    ]


# ---------------------------------------------------------------------------
# Unit tests
# ---------------------------------------------------------------------------

class Result:
    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.failures: List[str] = []

    def check(self, label: str, cond: bool, detail: str = ""):
        if cond:
            ok(label)
            self.passed += 1
        else:
            fail(f"{label}  ← {detail}")
            self.failed += 1
            self.failures.append(f"{label}: {detail}")


def run_unit_tests(cases: List[Dict]) -> Result:
    print(f"\n{BOLD}=== Unit Tests — field-level invariants ==={RESET}\n")
    r = Result()
    by = {c["symbol"]: c for c in cases}

    print(f"{CYAN}[dispatch trace]{RESET}")
    for sym, exp in [
        ("INVALIDATEDUSDT", "not_routed"),
        ("EXPIREDUSDT",     "not_routed"),
        ("REJECTEDUSDT",    "not_routed"),
        ("MAINSIGUSDT",     "MAIN_SIGNAL"),
        ("WATCHLISTUSDT",   "WATCHLIST"),
        ("NOSENDUSDT",      "NO_SEND"),
        ("PENDINGUSDT",     "not_evaluated"),
    ]:
        got = by[sym]["dispatch_action"]
        r.check(f"{sym:20} dispatch_action={exp}", got == exp, f"got={got}")

    for sym in ("INVALIDATEDUSDT", "EXPIREDUSDT", "REJECTEDUSDT"):
        got = by[sym]["dispatch_confidence_band"]
        r.check(f"{sym:20} dispatch_confidence_band=none", got == "none", f"got={got}")

    print(f"\n{CYAN}[semantic consistency]{RESET}")
    r.check("MAINSIGUSDT       semantic_consistency=Y",
            by["MAINSIGUSDT"]["semantic_consistency"] == "Y")
    r.check("SEMBROKENUSDT     semantic_consistency=N (missing confirmed_ts_ms)",
            by["SEMBROKENUSDT"]["semantic_consistency"] == "N",
            f"got={by['SEMBROKENUSDT']['semantic_consistency']}")
    r.check("SENTBROKENUSDT    semantic_consistency=N (sent but missing sent_ts_ms)",
            by["SENTBROKENUSDT"]["semantic_consistency"] == "N",
            f"got={by['SENTBROKENUSDT']['semantic_consistency']}")
    r.check("SENTBROKENUSDT    sent_ts_ms is empty",
            str(by["SENTBROKENUSDT"]["sent_ts_ms"]).strip() == "",
            f"got={by['SENTBROKENUSDT']['sent_ts_ms']}")

    print(f"\n{CYAN}[close semantics]{RESET}")
    r.check("NOTDUEUSDT        close_capture_basis=not_due_yet",
            by["NOTDUEUSDT"]["close_capture_basis"] == "not_due_yet")
    r.check("NOTDUEUSDT        closed_ts_ms empty — no fake timestamp",
            str(by["NOTDUEUSDT"]["closed_ts_ms"]).strip() in ("", "0"))
    r.check("FALLBACK4HUSDT    close_capture_basis=fallback_4h_snapshot",
            by["FALLBACK4HUSDT"]["close_capture_basis"] == "fallback_4h_snapshot",
            f"got={by['FALLBACK4HUSDT']['close_capture_basis']}")
    r.check("FALLBACK4HUSDT    close_reason present (not broken)",
            str(by["FALLBACK4HUSDT"]["close_reason"]).strip() != "",
            "empty — would be counted as broken close")
    r.check("CLOSEBRKNUSDT     close_reason empty (intentionally broken)",
            str(by["CLOSEBRKNUSDT"]["close_reason"]).strip() == "")
    r.check("CLOSEBRKNUSDT     status=INVALIDATED (terminal)",
            by["CLOSEBRKNUSDT"]["status"] == "INVALIDATED")

    print(f"\n{CYAN}[Bug 1 regression — closed_no_reason counter]{RESET}")
    _TERMINAL = {"INVALIDATED", "EXPIRED_WAIT"}
    should_count = [
        c for c in cases
        if str(c.get("status","")).upper() in _TERMINAL
        and str(c.get("closed_ts_ms","")).strip() not in ("", "0", "not_reached_yet")
        and not str(c.get("close_reason","")).strip()
    ]
    r.check("closed_no_reason count == 1 (only CLOSEBRKNUSDT)",
            len(should_count) == 1,
            f"counted {len(should_count)}: {[c['symbol'] for c in should_count]}")
    r.check("NOTDUEUSDT not counted as broken close",
            "NOTDUEUSDT" not in [c["symbol"] for c in should_count])
    r.check("FALLBACK4HUSDT not counted as broken close",
            "FALLBACK4HUSDT" not in [c["symbol"] for c in should_count])

    print(f"\n{CYAN}[outcome fields]{RESET}")
    r.check("WAITSHORTUSD      outcome_conclusion=WAIT_TOO_SHORT_CANDIDATE",
            by["WAITSHORTUSD"]["outcome_conclusion_code"] == "WAIT_TOO_SHORT_CANDIDATE")
    r.check("WAITSHORTUSD      regret_valid=Y",
            by["WAITSHORTUSD"]["regret_valid_YN"] == "Y")
    r.check("TRUECLOSEUSDT     outcome_conclusion=KILL_CORRECT",
            by["TRUECLOSEUSDT"]["outcome_conclusion_code"] == "KILL_CORRECT")

    return r


# ---------------------------------------------------------------------------
# DOCX helpers — extract specific cell values from tables
# ---------------------------------------------------------------------------

def _find_table_value(doc, section_title: str, row_label: str) -> Optional[str]:
    """
    Dual-mode value extractor:
    - If section renders as a TABLE (Section H): find first table after heading,
      return second cell of matching row.
    - If section renders as STATUS_LINE paragraphs (Section B): scan paragraphs
      after heading for row_label, extract value after last ': '.
    status_line() format: "✓  <label>: <value>"
    """
    from docx.oxml.ns import qn
    from docx.table import Table

    found = False
    para_texts = []  # collect paragraphs after section heading

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
        if tag == "p":
            text = "".join(r.text for r in block.iter(qn("w:t")))
            if section_title in text:
                found = True
                para_texts = []
                continue
            if found:
                para_texts.append(text)
        elif tag == "tbl" and found:
            # Try table-mode first (Section H style)
            tbl = Table(block, doc)
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if cells and row_label in cells[0]:
                    return cells[1] if len(cells) > 1 else None
            # Table found but label not in it — fall through to para_texts
            break

    # Para-mode fallback (Section B status_line style)
    for text in para_texts:
        if row_label in text and ": " in text:
            return text.rsplit(": ", 1)[-1].strip()
    return None


def _find_status_in_table(doc, section_title: str, row_label: str) -> Optional[str]:
    """Return last cell (status) of matched table row."""
    from docx.oxml.ns import qn
    from docx.table import Table

    found = False
    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
        if tag == "p":
            text = "".join(r.text for r in block.iter(qn("w:t")))
            if section_title in text:
                found = True
        elif tag == "tbl" and found:
            tbl = Table(block, doc)
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if cells and row_label in cells[0]:
                    return cells[-1] if cells else None
            return None
    return None


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Report integration test
# ---------------------------------------------------------------------------

def run_report_test(cases: List[Dict], workdir: Path, builder: Path) -> Result:
    print(f"\n{BOLD}=== Report Integration Test — DOCX build + value assertions ==={RESET}\n")
    r = Result()

    # Write CSV
    pending_path = workdir / f"pending_{DAY}.csv"
    fieldnames = list(cases[0].keys())
    with pending_path.open("w", newline="", encoding="utf-8") as f:
        csv.DictWriter(f, fieldnames=fieldnames).writeheader()
        csv.DictWriter(f, fieldnames=fieldnames).writerows(cases)
    # rewrite properly
    with pending_path.open("w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        w.writerows(cases)
    info(f"Wrote {len(cases)} cases → {pending_path.name}")

    ws      = workdir / "workspace"; ws.mkdir(exist_ok=True)
    out_dir = workdir / "out";       out_dir.mkdir(exist_ok=True)

    cmd = [
        sys.executable, str(builder),
        "--date", DAY,
        "--workspace", str(ws),
        "--pending",   str(pending_path),
        "--signals", "", "--results", "",
        "--snapshot-index", "",
        "--out-dir",   str(out_dir),
        "--tz", "Asia/Ho_Chi_Minh",
        "--fallback-hours", "4",
    ]
    info("Running builder...")
    result = subprocess.run(cmd, capture_output=True, text=True)

    r.check("Builder exits 0 (no crash)",
            result.returncode == 0,
            f"returncode={result.returncode}\n{result.stderr[:300]}")
    if result.returncode != 0:
        print(f"  stderr: {result.stderr[:500]}")
        return r

    docx_path = out_dir / f"daily_review_{DAY}.docx"
    r.check("DOCX file written", docx_path.exists())
    if not docx_path.exists():
        return r

    size_kb = docx_path.stat().st_size / 1024
    r.check(f"DOCX size > 20KB (got {size_kb:.0f}KB)", size_kb > 20,
            "too small — likely empty or crashed mid-render")

    try:
        from docx import Document
        doc = Document(str(docx_path))
        all_text = _all_text(doc)

        # ── Sections A-L ──
        print(f"\n{CYAN}[Sections A-L exist]{RESET}")
        for letter, title in [
            ("A", "Data Quality Gate"),
            ("B", "Semantic Health Panel"),
            ("C", "Confirm Path Integrity"),
            ("D", "Executive Decision Summary"),
            ("E", "Strategy Outcome Matrix"),
            ("F", "Missed Opportunity Board"),
            ("G", "Bad Loss Filter Board"),
            ("H", "Decision Trace Coverage"),
            ("I", "Human Review Queue"),
            ("J", "Case Registry Preview"),
            ("L", "Decision Footer"),
        ]:
            r.check(f"Section {letter} — {title}", title in all_text)

        # ── Section B: actual counter values ──
        print(f"\n{CYAN}[Section B — semantic counter values]{RESET}")

        closed_val = _find_table_value(doc, "Semantic Health Panel",
                                       "Closed rows missing close_reason")
        r.check(
            "Section B: closed_no_reason == 1 (only CLOSEBRKNUSDT)",
            closed_val == "1",
            f"got='{closed_val}'. >1 = Bug 1 regression. 0 = broken archetype not counted.")

        sent_val = _find_table_value(doc, "Semantic Health Panel",
                                     "Sent rows missing sent_ts_ms")
        r.check(
            "Section B: sent_no_ts == 1 (only SENTBROKENUSDT)",
            sent_val == "1",
            f"got='{sent_val}'")

        r.check(
            "Section B: health banner = ISSUES FOUND (real broken cases exist)",
            "ISSUES FOUND" in all_text,
            "expected ISSUES FOUND — broken rows exist in dataset")

        # ── Section C: not_reached_yet ──
        r.check(
            "Section C: not_reached_yet rendered (NOTDUEUSDT close semantic)",
            "not_reached_yet" in all_text,
            "not_reached_yet not found — close semantic rendering broken")

        # ── Section H: dispatch coverage ──
        print(f"\n{CYAN}[Section H — dispatch coverage]{RESET}")

        dispatch_pop = _find_table_value(doc, "Decision Trace Coverage",
                                         "dispatch_action")
        r.check(
            "Section H: dispatch_action populated count > 0 (Phase 5A fix)",
            dispatch_pop is not None and dispatch_pop.strip() not in ("0", ""),
            f"got='{dispatch_pop}' — still 0 means Phase 5A fix not working")

        dispatch_status = _find_status_in_table(doc, "Decision Trace Coverage",
                                                "dispatch_action")
        r.check(
            "Section H: dispatch status is explicit (N/A, OK, or GAPS)",
            dispatch_status in ("N/A", "GAPS", "OK"),
            f"got='{dispatch_status}' — blank or unexpected value")

        # ── Footer / next_action ──
        print(f"\n{CYAN}[Footer — next_action alignment]{RESET}")

        r.check(
            "Footer: next_action = Fix semantic health first",
            "Fix semantic health first" in all_text,
            "expected semantic warning — broken rows present in dataset")

        r.check(
            "Footer: NOT 'Data clean' when broken rows present",
            "Data clean — no single highest-priority action today." not in all_text,
            "Footer says Data clean but broken rows exist — next_action logic wrong")

        info(f"DOCX: {docx_path}")

    except Exception as e:
        fail(f"DOCX parse error: {e}")
        import traceback; traceback.print_exc()
        r.failed += 1
        r.failures.append(str(e))

    return r


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--report-only", action="store_true")
    parser.add_argument("--no-report",   action="store_true")
    parser.add_argument("--keep",        action="store_true")
    args = parser.parse_args()

    builder = Path("build_daily_review_pack.py")
    if not builder.exists():
        print(f"{RED}ERROR: build_daily_review_pack.py not found.{RESET}")
        print("Run from the same directory as the builder.")
        sys.exit(1)

    print(f"\n{BOLD}{'='*60}")
    print("  Simulation Regression Suite")
    print(f"{'='*60}{RESET}")
    print(f"  Builder : {builder.resolve()}")
    print(f"  Date    : {DAY}")
    print(f"  Cases   : 14 synthetic archetypes\n")

    cases = build_cases()
    all_results: List[Result] = []

    if not args.report_only:
        all_results.append(run_unit_tests(cases))

    if not args.no_report:
        workdir = Path(tempfile.mkdtemp(prefix="sim_test_"))
        info(f"Temp dir: {workdir}")
        try:
            all_results.append(run_report_test(cases, workdir, builder))
        finally:
            if not args.keep:
                shutil.rmtree(workdir, ignore_errors=True)
            else:
                info(f"Kept: {workdir}")

    total_pass   = sum(r.passed  for r in all_results)
    total_fail   = sum(r.failed  for r in all_results)
    all_failures = [f for r in all_results for f in r.failures]

    print(f"\n{BOLD}{'='*60}  SUMMARY  {'='*60}{RESET}")
    print(f"  Passed : {GREEN}{total_pass}{RESET}")
    print(f"  Failed : {(RED if total_fail else GREEN)}{total_fail}{RESET}")

    if all_failures:
        print(f"\n{RED}Failures:{RESET}")
        for f in all_failures:
            print(f"  • {f}")
        print()
        sys.exit(1)
    else:
        print(f"\n{GREEN}{BOLD}  ALL TESTS PASSED ✓{RESET}\n")
        sys.exit(0)


if __name__ == "__main__":
    main()
